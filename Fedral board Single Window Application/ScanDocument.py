#Import the required Libraries
# from thanksWindow import *
from tkinter import *

# from ConfirmDocument import * for tommorow ..
import tkinter as tk
from PIL import Image, ImageTk
import os
# from thanksWindow import *
from datetime import datetime
import shutil
import pypyodbc as pyodbc
import pandas as pd
import os
from openpyxl import load_workbook
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert
from pdf2image import convert_from_path
from PIL import Image,ImageTk
from pdf2image import convert_from_path
from docx2pdf import convert
from PIL import Image,ImageTk
import time
import pywintypes
import win32api
import docx
from docx.shared import Inches
import io


check = 1
# connecting with the payments database for getting the reg_no on the refference of challan no
import mysql.connector
from mysql.connector import Error

try:
    connection = mysql.connector.connect(host='192.168.100.2',
                                         database='fee_challans',
                                        #  user='yaseen',
                                        #  password='Smile@123asdf'
                                         user='yaseen2',
                                         password='fVum*.NODLS]w_6F'
    )
                                        
    if connection.is_connected():
        db_Info = connection.get_server_info()
        print("Connected to MySQL Server version ", db_Info)
        cursor_P = connection.cursor()
        # cursor_P.execute("SELECT reg_no FROM `payments` WHERE challan_no = 771666175600 ")
        # record = cursor_P.fetchone()[0]
        # print(record)

except Error as e:
    print("Error while connecting to MySQL", e)

#  connection with DB ssc
try:
    cnxn_str = ("Driver={SQL Server Native Client 11.0};"
            "Server=DBSERVER;"
            "Database=ssc;"
            "UID=kiosk01;"
            "PWD=123;")

    cnxn = pyodbc.connect(cnxn_str)


 
except:
    
    def OpenScanDocumentFile():
        os.system("python ScanDocument.py")
    window = Tk()
    window.geometry('1600x1600')
    window.title("DB connection Failed")
    Error = tk.Label(window, text="Error", font=("Microsoft JhengHei", 18))
    Error.place(relx = 0.4,
                    rely = 0.2,
                    anchor = 'center') 
    Database_error = tk.Label(window, text="Database Connection Failed Please Connect LAN Cable", font=("Microsoft JhengHei", 16))
    Database_error.place(relx = 0.4,
                    rely = 0.3,
                    anchor = 'center') 
    my_button_retry = Button(text="Retry", relief="raised", font=("Times New Roman", 16),  command=lambda : [ OpenScanDocumentFile(), ] )
    my_button_retry.place(relx = 0.4,
                    rely = 0.4,
                    anchor = 'center',
                    )  
    window.mainloop()
    
# cnxn_str = ("Driver={SQL Server Native Client 11.0};"
#             "Server=DBSERVER;"
#             "Database=ssc;"
#             "UID=kiosk01;"
#             "PWD=123;")

# cnxn = pyodbc.connect(cnxn_str)



# connection with db ssc ledger 
try:
    cnxn_str_2 = ("Driver={SQL Server Native Client 11.0};"
                "Server=DBSERVER;"
                "Database=ssc_LEDGER;"
                "UID=kiosk01;"
                "PWD=123;")

    cnxn_2 = pyodbc.connect(cnxn_str_2)
except:    
    def OpenScanDocumentFile():
        os.system("python ScanDocument.py")
    window = Tk()
    window.geometry('1600x1600')
    window.title("DB Connection Failed")
    Error = tk.Label(window, text="Error", font=("Microsoft JhengHei", 18))
    Error.place(relx = 0.4,
                    rely = 0.2,
                    anchor = 'center') 
    my_button_retry = Button(text="Retry", relief="raised", font=("Times New Roman", 16),  command=lambda : [window.destroy(), OpenScanDocumentFile()] )
    my_button_retry.place(relx = 0.4,
                    rely = 0.4,
                    anchor = 'center',
                    )
    Database_error = tk.Label(window, text="Database Connection Failed Please Connect LAN Cable", font=("Microsoft JhengHei", 16))
    Database_error.place(relx = 0.4,
                    rely = 0.3,
                    anchor = 'center')  
    window.mainloop()
    
    

try:
    cnxn_str_3 = ("Driver={SQL Server Native Client 11.0};"
            "Server=DBSERVER;"
            "Database=hssc;"
            "UID=kiosk01;"
            "PWD=123;")

    cnxn_3 = pyodbc.connect(cnxn_str_3)


 
except:
    
    def OpenScanDocumentFile():
        os.system("python ScanDocument.py")
    window = Tk()
    window.geometry('1600x1600')
    window.title("DB connection Failed")
    Error = tk.Label(window, text="Error", font=("Microsoft JhengHei", 18))
    Error.place(relx = 0.4,
                    rely = 0.2,
                    anchor = 'center') 
    Database_error = tk.Label(window, text="Database Connection Failed Please Connect LAN Cable", font=("Microsoft JhengHei", 16))
    Database_error.place(relx = 0.4,
                    rely = 0.3,
                    anchor = 'center') 
    my_button_retry = Button(text="Retry", relief="raised", font=("Times New Roman", 16),  command=lambda : [ OpenScanDocumentFile(), ] )
    my_button_retry.place(relx = 0.4,
                    rely = 0.4,
                    anchor = 'center',
                    )  
    window.mainloop()
    
    
try:
    cnxn_str_4 = ("Driver={SQL Server Native Client 11.0};"
            "Server=DBSERVER;"
            "Database=Hssc_LEDGER;"
            "UID=kiosk01;"
            "PWD=123;")

    cnxn_4 = pyodbc.connect(cnxn_str_4)


 
except:
    
    def OpenScanDocumentFile():
        os.system("python ScanDocument.py")
    window = Tk()
    window.geometry('1600x1600')
    window.title("DB connection Failed")
    Error = tk.Label(window, text="Error", font=("Microsoft JhengHei", 18))
    Error.place(relx = 0.4,
                    rely = 0.2,
                    anchor = 'center') 
    Database_error = tk.Label(window, text="Database Connection Failed Please Connect LAN Cable", font=("Microsoft JhengHei", 16))
    Database_error.place(relx = 0.4,
                    rely = 0.3,
                    anchor = 'center') 
    my_button_retry = Button(text="Retry", relief="raised", font=("Times New Roman", 16),  command=lambda : [ OpenScanDocumentFile(), ] )
    my_button_retry.place(relx = 0.4,
                    rely = 0.4,
                    anchor = 'center',
                    )  
    window.mainloop()

try:
    cnxn_str_5 = ("Driver={SQL Server Native Client 11.0};"
            "Server=192.168.100.1\SQL2005;"
            "Database=Accounts;"
            "UID=challan_verify;"
            "PWD=verification;")

    cnxn_5 = pyodbc.connect(cnxn_str_5)


 
except:
    
    def OpenScanDocumentFile():
        os.system("python ScanDocument.py")
    window = Tk()
    window.geometry('1600x1600')
    window.title("DB connection Failed")
    Error = tk.Label(window, text="Error", font=("Microsoft JhengHei", 18))
    Error.place(relx = 0.4,
                    rely = 0.2,
                    anchor = 'center') 
    Database_error = tk.Label(window, text="Database Connection Failed Please Connect LAN Cable", font=("Microsoft JhengHei", 16))
    Database_error.place(relx = 0.4,
                    rely = 0.3,
                    anchor = 'center') 
    my_button_retry = Button(text="Retry", relief="raised", font=("Times New Roman", 16),  command=lambda : [ OpenScanDocumentFile(), ] )
    my_button_retry.place(relx = 0.4,
                    rely = 0.4,
                    anchor = 'center',
                    )  
    window.mainloop()

cursor_2 = cnxn_2.cursor()

cursor_3 = cnxn_3.cursor()

cursor_4 = cnxn_4.cursor()

cursor = cnxn.cursor()

cursor_5 =  cnxn_5.cursor()




# Current date.
today = datetime.today()
Dated = today.date()
 
 
# import  datetime

# # Current date.
# current_year  = datetime.date.today()
# # Dated = today.date()
# year = current_year.strftime("%Y")
# print(year)
nbr = '9010110004'
               
Name_display = ''' SELECT name FROM ZReg WHERE reg_no =?'''

cursor.execute(Name_display, [nbr])
    
Name_display_O = cursor.fetchone()[0]

import pypyodbc
import base64
from base64 import * 

# connection = pypyodbc.connect('Driver=SQL Server;'
#                             'Server=DESKTOP-MSSQLSERVER;'
#                             'Database=Test;'
#                             'Trusted_Connection=yes;'
#                             )

# cursor = connection.cursor()

# a = 'bob@bob.com'
# b = 'mack jones'
# filename = 'bookingsuccessful.PNG'


# image = open(filename, 'rb')
# image_read = image.read()
# image_64_encode = base64.encodebytes(image_read)


# image.close()

# SQLCommand = ("INSERT INTO Validation(email, myname, photo) VALUES(?,?,?)")
# Values = [a,b,image_64_encode]
# cursor.execute(SQLCommand, Values)
# connection.commit()

# photo_path = r'C:\Users\hp\Desktop\Fedral board project\Fedral Board\Fedral board Project\Fedral board Single Window Application' + '\\'
# retrieved_bytes = cursor_2.execute("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = '2029811146' ")
# with open(photo_path + 'new.jpg', 'wb') as new_jpg:
#     new_jpg.write(retrieved_bytes)
# print(f'{len(retrieved_bytes)} bytes retrieved and written to new file')

  
# Create an instance of a word document
# SQLCommand = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = '2029811146' ")
# doc = docx.Document()
  
# # Add a Title to the document
# # doc.add_heading('GeeksForGeeks', 0)

# cursor_2.execute(SQLCommand)
# data = cursor_2.fetchone()[0]


# doc.add_picture(io.BytesIO(data))
  
# # Image with defined size
# doc.add_heading('Image with Defined Size:', 3)
# # doc.add_picture('sir.png', width=Inches(2), height=Inches(2))
  
# # Now save the document to a location
# doc.save('gfg.docx')





# data = bytes(data.strip("\n"), 'utf-8')

# image_64_decode = base64.decodebytes(data)
# image_result = open('testfile.gif', 'wb')
# image_result.write(image_64_decode)
# image_result.close()

# connection.close()


# Condition to select the Certificate

# select reg_no from challan where ch_no = '981273123'

# certificate = select certificate from challan where ch_no = '89371293'

# if(certifiacte == "MIGRATION CERTIFICATE")

# # For Migration Certificate        

# M_C_Name = cursor.execute("SELECT name from ZReg WHERE reg_no = '9010110004'")

# Name_M_C = M_C_Name.fetchone()[0]

# M_C_DOB = cursor.execute("SELECT dob from ZReg WHERE reg_no = '9010110004'")

# DOB_M_C = M_C_DOB.fetchone()[0]

# M_C_fname = cursor.execute("Select fname from ZReg WHERE reg_no = '9010110004'")

# fname_M_C = M_C_fname.fetchone()[0]

# M_C_reg_no = cursor.execute("Select reg_no from ZReg where fname = 'KARIM KHAN MARWAT'")

# reg_no_M_C = M_C_reg_no.fetchone()[0]

# M_C_roll_no = cursor_2.execute("select roll_no from ZLedgerII where reg_no = '9010110004'")

# roll_no_M_C = M_C_roll_no.fetchone()[0]

# M_C_Year = cursor.execute("Select year from ZReg WHERE reg_no = '9010110004'")

# Year_M_C = M_C_Year.fetchone()[0]

# M_C_Institution = cursor.execute("select inst_desc from ZReg where reg_no = '9010110004'")

# Institution_M_C = M_C_Institution.fetchone()[0]


# /////// Completed Migration Certificate

# # for Result Cancelation Certificate

# R_C_C_Name = cursor.execute("SELECT name from ZReg WHERE reg_no = '9010110004'")

# Name_R_C_C = R_C_C_Name.fetchone()[0]

# R_C_C_fname = cursor.execute("Select fname from ZReg WHERE reg_no = '9010110004'")

# fname_R_C_C = R_C_C_fname.fetchone()[0]

# R_C_C_roll_no = cursor_2.execute("select roll_no from ZLedgerII where reg_no = '9010110004'")

# roll_no_R_C_C = R_C_C_roll_no.fetchone()[0]

# R_C_C_reg_no = cursor.execute("Select reg_no from ZReg where fname = 'KARIM KHAN MARWAT'")

# reg_no_R_C_C = R_C_C_reg_no.fetchone()[0]

# Completed Result Cancelation Certificate

# For Application form for migration request

# AMR_Name = cursor.execute("SELECT name from ZReg WHERE reg_no = '9010110004'")

# Name_AMR = AMR_Name.fetchone()[0]

# AMR_fname = cursor.execute("Select fname from ZReg WHERE reg_no = '9010110004'")

# fname_AMR = AMR_fname.fetchone()[0]

# AMR_roll_no = cursor_2.execute("select roll_no from ZLedgerII where reg_no = '9010110004'")

# roll_no_AMR = AMR_roll_no.fetchone()[0]

# AMR_reg_no = cursor.execute("Select reg_no from ZReg where fname = 'KARIM KHAN MARWAT'")

# reg_no_AMR = AMR_reg_no.fetchone()[0]

# AMR_Year = cursor.execute("Select year from ZReg WHERE reg_no = '9010110004'")

# Year_AMR = AMR_Year.fetchone()[0]

# cursor_P.execute("SELECT challan_no FROM payments WHERE head_code = 0300313")
# a = cursor_P.fetchone()[0]
# print(a)





# gloabal variables used in this system



# completed Application form for migration certificate

# For making Entry only in integer
# class Lotfi(tk.Entry):
#     def __init__(self, master=None, **kwargs):
#         self.var = tk.StringVar()
#         tk.Entry.__init__(self, master, textvariable=self.var, **kwargs)
#         self.old_value = '' 
#         self.var.trace('w', self.check)
#         self.get, self.set = self.var.get, self.var.set

#     def check(self, *args):
#         if self.get().isdigit(): 
#             # the current value is only digits; allow this
#             self.old_value = self.get() 
#         else:
#             # there's non-digit characters in the input; reject this 
#             self.set(self.old_value)

def disable_event():
       pass
   
   
root = Tk();


root.title("Fedral Board")

# root.protocol("WM_DELETE_WINDOW", disable_event)


#dimensions
# canvas = tk.Canvas(root, width=1600, height=1600)
# canvas.grid(columnspan=3, rowspan=3)



#logo 
logo = Image.open('sir.png')
logo = ImageTk.PhotoImage(logo)
logo_label = tk.Label(image=logo)
logo_label.image = logo
logo_label.grid(column=1, row=0)
root.attributes('-fullscreen',True)






# #instructions
# instructions = tk.Label(root, text="Please scan barcode From your Chalan form",bg='#add8e6'
# , relief="raised",font=("Times New Roman", 22))
# instructions.place(relx = 0.5,
#                 rely = 0.3,
#                 anchor = 'center')
# instructions2 = tk.Label(root, text="Or",bg='#add8e6', relief="raised", font=("Times New Roman", 22))
# instructions2.place(relx = 0.5,
#                 rely = 0.4,
#                 anchor = 'center')

# instructions3 = tk.Label(root, text="Type chalan number",bg='#add8e6',relief="raised", font=("Times New Roman", 22))
# instructions3.place(relx = 0.5,
#                 rely = 0.5,
#                 anchor = 'center')

# From_entry = Lotfi(root, width=25)



# From_entry.place(relx = 0.7,
#                 rely = 0.6,
#                 anchor = 'center',
#                 width=200,
#                 height=30
#                 )



# For closing the program
def Closing_Window():
    
    win = Tk()
    win.geometry('600x300')
    win.title("Closing Window")
    instructions_closing = tk.Label(win, text="NOT FOR STUDENTS", font=("Microsoft JhengHei", 14))
    instructions_closing.place(relx = 0.3,
                    rely = 0.3,
                    anchor = 'center')
    
    pass_var=Entry(win, show="*")
    pass_var.place(relx = 0.3,
                    rely = 0.4,
                    anchor = 'center',
                    width=200,
                    height=30
                    )
    
    def close_root():
        pass_value = pass_var.get()
        if(pass_value == "close"):
            root.destroy()
            win.destroy()
           
       
    
    #Create a button to close the window
    btnclose = tk.Button(win, text ="Click here to Close",command=close_root)
    btnclose.place(relx = 0.3,
                    rely = 0.5,
                    anchor = 'center',
                
                    )




    
    win.mainloop()  
btn = Button(root, text="X", command = Closing_Window,bg='red')
btn.place(relx = 0.9,
                rely = 0.1,
                anchor = 'center',
                width=50,
                height=30
                
                ) 

            # global variable 
# showing all data in display 


# end 


# function clear button

# def globala():
#     global exp


# globala()
exp = " "          # global variable 
# showing all data in display 

def press(num):
    global exp
    exp=exp + str(num)
    equation.set(exp)
# end 



# function clear button

def clear():
    global exp
    exp = " "
    equation.set(exp)

# end 


# Enter Button Work Next line Function

def action():
    exp = " Next Line : "
    equation.set(exp)

# end function coding




# action()




# # Size rootow size
# root.geometry('1010x250')         # normal size
# root.maxsize(width=1010, height=250)      # maximum size
# root.minsize(width= 1010 , height = 250)     # minimum size
# # end rootow size

root.configure(bg = 'green')    #  add background color

# entry box
equation = StringVar()
from_entry = Entry(root,textvariable = equation, font=('Helvetica bold', 23))
from_entry.place(relx=0.5, rely=0.5, width=400, height=50, )
# end entry box

# add all button line wise 

# First Line Button

q = Button(root,text = '1' , width = 6, font=('Helvetica bold', 18), command = lambda : press('1'))
q.place(x=840, y=570, width=50, height=50)

w = Button(root,text = '2' , width = 6,font=('Helvetica bold', 18), command = lambda : press('2'))
w.place(x=890, y=570,width=50, height=50)

E = Button(root,text = '3' , width = 6,font=('Helvetica bold', 18), command = lambda : press('3'))
E.place(x=940, y=570,width=50, height=50)


A = Button(root,text = '4' , width = 6, font=('Helvetica bold', 18),command = lambda : press('4'))
A.place(x=840, y=620,width=50, height=50)

clear = Button(root,text = 'Clear' , width = 6,command = clear)
clear.place(x=1000, y=635,width=50, height=50)
# clear.grid(row = 1 , column = 5, ipadx = 6 , ipady = 10)



S = Button(root,text = '5' , width = 6, font=('Helvetica bold', 18),command = lambda : press('5'))
S.place(x=890, y=620,width=50, height=50)

D = Button(root,text = '6' , width = 6,font=('Helvetica bold', 18), command = lambda : press('6'))
D.place(x=940, y=620,width=50, height=50)


Z = Button(root,text = '7' , width = 6,font=('Helvetica bold', 18), command = lambda : press('7'))
Z.place(x=840, y=670,width=50, height=50)


X = Button(root,text = '8' , width = 6,font=('Helvetica bold', 18), command = lambda : press('8'))
X.place(x=890, y=670,width=50, height=50)


C = Button(root,text = '9' , width = 6,font=('Helvetica bold', 18), command = lambda : press('9'))
C.place(x=940, y=670,width=50, height=50)

C = Button(root,text = '0' , width = 6,font=('Helvetica bold', 18), command = lambda : press('0'))
C.place(x=840, y=720,width=120, height=45)

# C = Button(root,text = 'DEL' , width = 6,font=('Helvetica bold', 18), command = lambda : press('0'))
# C.place(x=1000, y=590,width=120, height=45)


def DEl():
    pass

def backspace():
    if Dis_entry == 0:
            Dis_entry
            


#   function for getting the challan number from the textbox s
# def challan_Entry():
    
    
#     challan_No = From_entry.get()

#     nbr = challan_No
            
#     command = ''' SELECT reg_no FROM payments WHERE challan_no =%s'''

#     cursor_P.execute(command, [nbr])
#     global reg_no_
    
#     reg_no_ = cursor_P.fetchone()[0]
    
#     # # for Result Cancelation Certificate

#     R_C_C_Name = cursor.execute("SELECT name from ZReg WHERE reg_no = '9010110004'")

#     Name_R_C_C = R_C_C_Name.fetchone()[0]

#     R_C_C_fname = cursor.execute("Select fname from ZReg WHERE reg_no = '9010110004'")

#     fname_R_C_C = R_C_C_fname.fetchone()[0]

#     R_C_C_roll_no = cursor_2.execute("select roll_no from ZLedgerII where reg_no = '9010110004'")

#     roll_no_R_C_C = R_C_C_roll_no.fetchone()[0]

#     R_C_C_reg_no = cursor.execute("Select reg_no from ZReg where fname = 'KARIM KHAN MARWAT'")

#     reg_no_R_C_C = R_C_C_reg_no.fetchone()[0]

# Completed Result Cancelation Certificate

    
    
    
    # command2 = ''' SELECT reg_no FROM payments WHERE challan_no =%s'''

    
    
    
    # print(reg_no_)
    

    
    # print(data)
        

# from one screen to another

my_button = Button(text="Check" ,relief="raised",font=("Times New Roman", 22),  command=lambda : [  Making_document() ] )
my_button.place(x = 900,
                y = 500, 
                anchor = 'center',
                width=300,
                height=86
                )


               



def Making_document():
    
    
    
    
    # try:
    #     challan_No = from_entry.get()

    #     nbr = challan_No
                
    #     command = ''' select verified from income where bank_chalan_number = ? '''


    #     cursor_5.execute(command, [nbr])
    #     pay_or_not = cursor_5.fetchone()[0]
    #     print(pay_or_not)
    #     if(pay_or_not== "1"):
    #         root = Tk()
    #         root.title("ALready Taken")
    #         root.geometry("600x400")

    #         instructions1_ = Label(root, text="You Already ", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #         instructions1_.place(relx = 0.4,
    #                             rely = 0.2,
    #                             anchor = 'center')

    #         instructions2_ = Label(root, text="Taken", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #         instructions2_.place(relx = 0.4,
    #                             rely = 0.3,
    #                             anchor = 'center')

    #         instructions3_ = Label(root, text="This servise", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #         instructions3_.place(relx = 0.4,
    #                             rely = 0.4,
    #                             anchor = 'center')

    #         root.mainloop()
    # except:
    try:
        challan_No = from_entry.get()
        

        nbr = challan_No
                
        command = ''' SELECT reg_no FROM payments WHERE challan_no =%s'''
        

        cursor_P.execute(command, [nbr])

        
        reg_no_ = cursor_P.fetchone()[0]
        print(reg_no_)
        
        curr_year = ''' select year from ZLedger where  reg_no = ?'''

        cursor_2.execute(curr_year, [reg_no_])

        year_curr = cursor_2.fetchone()[0]
    
    except:
        root = Tk()
        root.title("Challan error")
        root.geometry("600x400")

        instructions1_ = Label(root, text="Your Challan Number ", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
        instructions1_.place(relx = 0.4,
                            rely = 0.2,
                            anchor = 'center')

        instructions2_ = Label(root, text="is incorrect", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
        instructions2_.place(relx = 0.4,
                            rely = 0.3,
                            anchor = 'center')

        instructions3_ = Label(root, text="please check your Challan number again", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
        instructions3_.place(relx = 0.4,
                            rely = 0.4,
                            anchor = 'center')

        root.mainloop()
    # if(pay_or_not):
    #     pass
    # else:
    #     root = Tk()
    #     root.title("Not paid")
    #     root.geometry("600x400")

    #     instructions1_ = Label(root, text="Service ", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #     instructions1_.place(relx = 0.4,
    #                         rely = 0.2,
    #                         anchor = 'center')

    #     instructions2_ = Label(root, text="Not", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #     instructions2_.place(relx = 0.4,
    #                         rely = 0.3,
    #                         anchor = 'center')

    #     instructions3_ = Label(root, text="Paid", bg='#add8e6', relief="raised", font=("Times New Roman", 22))
    #     instructions3_.place(relx = 0.4,
    #                         rely = 0.4,
    #                         anchor = 'center')

    #     root.mainloop()
    
    
    certificate_name = ''' SELECT head_code FROM `payments` WHERE challan_no  =%s'''
    cursor_P.execute(certificate_name, [nbr])
    certificate_name_code =  cursor_P.fetchone()[0]
    print(certificate_name_code)
    
    
    payment_status = ''' SELECT payment_status FROM `payments` WHERE challan_no =%s'''
    cursor_P.execute(payment_status, [nbr])
    status_P =  cursor_P.fetchone()[0]
    
    
    
    
    
    
    
    
        # Open files
    main_path = r"C:\Users\hp\Desktop\Fedral board project\Fedral Board\Fedral board Project\Fedral board Single Window Application"
    
        #Condition for selecting the template file which print 
    
    if certificate_name_code == '02000250':
        
    
        # # For Application form for migration request

        # AMR_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Name, [reg_no_])
        

        # Name_AMR = cursor.fetchone()[0]

        # AMR_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_fname, [reg_no_])

        # fname_AMR = cursor.fetchone()[0]

        # AMR_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        # cursor_2.execute(AMR_roll_no, [reg_no_])

        # roll_no_AMR = cursor_2.fetchone()[0]

        # AMR_reg_no = '''Select reg_no from ZReg where fname = ?'''
        # cursor.execute(AMR_reg_no, [fname_AMR])

        # reg_no_AMR = cursor.fetchone()[0]

        # AMR_Year = '''Select year from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Year, [reg_no_])

        # Year_AMR = cursor.fetchone()[0]
        
        # # For Migration Certificate        

        M_C_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor.execute(M_C_Name, [reg_no_])

        Name_M_C = cursor.fetchone()[0]
    
        M_C_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        cursor.execute(M_C_DOB, [reg_no_])

        DOB_M_C = cursor.fetchone()[0]

        M_C_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor.execute(M_C_fname, [reg_no_])

        fname_M_C = cursor.fetchone()[0]

        M_C_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor.execute(M_C_reg_no, [fname_M_C])

        reg_no_M_C = cursor.fetchone()[0]

        M_C_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        cursor_2.execute(M_C_roll_no, [reg_no_])

        roll_no_M_C = cursor_2.fetchone()[0]

        M_C_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor.execute(M_C_Year, [reg_no_])

        Year_M_C = cursor.fetchone()[0]

        M_C_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor.execute(M_C_Institution, [reg_no_])

        Institution_M_C = cursor.fetchone()[0]
        
        template_path = os.path.join(main_path, 'MIGRATION CERTIFIACTE_templ.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {'Candidate_name' : None,
                    'Dated' : None,
                    'DOB': None,
                    'Father_name': None,
                    'Registration_no' : None,
                    'Roll_no' : None,
                    'Examination': None,
                    'Session': None,
                    'year': None,
                    'Status': None,
                    'Institution': None
                    
                    }

    
    

        to_fill_in['Candidate_name'] = Name_M_C
        to_fill_in['DOB'] = DOB_M_C
        to_fill_in['Dated'] = Dated
        to_fill_in['Father_name'] = fname_M_C
        to_fill_in['Registration_no'] = reg_no_M_C
        to_fill_in['Roll_no'] = roll_no_M_C
        to_fill_in['Examination'] = "Fedral board"
        to_fill_in['Session'] = "Final Session"
        to_fill_in['year'] =  Year_M_C
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_M_C
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'MIGRATION CERTIFIACTE.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Done with MIGRATION CERTIFIACTE.docx")
        doc = Document('MIGRATION CERTIFIACTE.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_2.execute(data, [reg_no_])
        data = cursor_2.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('MIGRATION CERTIFIACTE.docx')

        
        convert("MIGRATION CERTIFIACTE.docx", "MIGRATION CERTIFIACTE.pdf")
        
        images = convert_from_path("MIGRATION CERTIFIACTE.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'MIGRATION CERTIFIACTE'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()
        wind.attributes('-fullscreen',True)
    

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        
                    
        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        

    

        #Load an image in the script
        img= (Image.open("MIGRATION CERTIFIACTE.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),OpenScanDocumentFile() ])
        btn_.place(x = 45,
                y=500,
                width=40,
                height=40
        
                
                ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile()])
        btn2_.place(x = 130,
                y=500,
                width=40,
                height=40
                
                )
        wind.mainloop()
    
    # completed Migration certificate Filling 
    
    if certificate_name_code == '03000350':
            
    
        # # For Application form for migration request

        # AMR_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Name, [reg_no_])
        

        # Name_AMR = cursor.fetchone()[0]

        # AMR_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_fname, [reg_no_])

        # fname_AMR = cursor.fetchone()[0]

        # AMR_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        # cursor_2.execute(AMR_roll_no, [reg_no_])

        # roll_no_AMR = cursor_2.fetchone()[0]

        # AMR_reg_no = '''Select reg_no from ZReg where fname = ?'''
        # cursor.execute(AMR_reg_no, [fname_AMR])

        # reg_no_AMR = cursor.fetchone()[0]

        # AMR_Year = '''Select year from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Year, [reg_no_])

        # Year_AMR = cursor.fetchone()[0]
        
        # # For Migration Certificate        

        M_C_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor_3.execute(M_C_Name, [reg_no_])

        Name_M_C = cursor_3.fetchone()[0]
    
        # M_C_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        # cursor_3.execute(M_C_DOB, [reg_no_])

        # DOB_M_C = cursor_3.fetchone()[0]

        M_C_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor_3.execute(M_C_fname, [reg_no_])

        fname_M_C = cursor_3.fetchone()[0]

        M_C_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor_3.execute(M_C_reg_no, [fname_M_C])

        reg_no_M_C = cursor_3.fetchone()[0]

        M_C_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        cursor_4.execute(M_C_roll_no, [reg_no_])

        roll_no_M_C = cursor_4.fetchone()[0]

        M_C_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor_3.execute(M_C_Year, [reg_no_])

        Year_M_C = cursor_3.fetchone()[0]

        M_C_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor_3.execute(M_C_Institution, [reg_no_])

        Institution_M_C = cursor_3.fetchone()[0]
        
        template_path = os.path.join(main_path, 'MIGRATION CERTIFIACTE_templ.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {'Candidate_name' : None,
                    'Dated' : None,
                    # 'DOB': None,
                    'Father_name': None,
                    'Registration_no' : None,
                    'Roll_no' : None,
                    'Examination': None,
                    'Session': None,
                    'year': None,
                    'Status': None,
                    'Institution': None
                    
                    }

    
    

        to_fill_in['Candidate_name'] = Name_M_C
        # to_fill_in['DOB'] = DOB_M_C
        to_fill_in['Dated'] = Dated
        to_fill_in['Father_name'] = fname_M_C
        to_fill_in['Registration_no'] = reg_no_M_C
        to_fill_in['Roll_no'] = roll_no_M_C
        to_fill_in['Examination'] = "Fedral board"
        to_fill_in['Session'] = "Final Session"
        to_fill_in['year'] =  Year_M_C
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_M_C
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'MIGRATION CERTIFIACTE.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Done with MIGRATION CERTIFIACTE.docx")
        doc = Document('MIGRATION CERTIFIACTE.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_4.execute(data, [reg_no_])
        data = cursor_4.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('MIGRATION CERTIFIACTE.docx')
        
        convert("MIGRATION CERTIFIACTE.docx", "MIGRATION CERTIFIACTE.pdf")
        
        images = convert_from_path("MIGRATION CERTIFIACTE.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'MIGRATION CERTIFIACTE'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        
        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()


        

        #Load an image in the script
        img= (Image.open("MIGRATION CERTIFIACTE.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)
        
        
        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(), OpenScanDocumentFile() ])
        btn_.place(x = 45,
                y=500,
                width=40,
                height=40
                
                        
                        ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile()])
        btn2_.place(x = 130,
                y=500,
                width=40,
                height=40
                        
                )


    
        
        wind.mainloop()
    
    # completed Migration certificate Filling 
    
    
    
    
    
    
    
    
    # # for Result Cancelation Certificate
    
    if certificate_name_code == '02000213':
        

        R_C_C_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_C_Name, [reg_no_])

        Name_R_C_C = cursor.fetchone()[0]

        R_C_C_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_C_fname, [reg_no_])
        

        fname_R_C_C = cursor.fetchone()[0]

        R_C_C_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        cursor_2.execute(R_C_C_roll_no, [reg_no_])

        roll_no_R_C_C = cursor_2.fetchone()[0]

        R_C_C_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor.execute(R_C_C_reg_no, [fname_R_C_C])

        reg_no_R_C_C = cursor.fetchone()[0]
        
        template_path = os.path.join(main_path, 'RESULT CANCELLATION CERTIFICATE_templ.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        # workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        # worksheet = workbook["Input"]

        to_fill_in = {'Candidate_name' : None,
                    'Father_name': None,
                    'Roll_no' : None,
                    'Registration_no' : None,
                    
                    }

        to_fill_in['Candidate_name'] = Name_R_C_C
        to_fill_in['Dated'] = Dated
        to_fill_in['Father_name'] = fname_R_C_C
        to_fill_in['Registration_no'] = reg_no_R_C_C
        to_fill_in['Roll_no'] = roll_no_R_C_C
        
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'RESULT CANCELLATION CERTIFICATE.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Done with RESULT CANCELLATION CERTIFICATE.docx")
        

        
        # Converting docx present in the same folder
        # as the python file
        # convert("RESULT CANCELLATION CERTIFICATE.docx")
        
        # Converting docx specifying both the input
        # and output paths
        convert("RESULT CANCELLATION CERTIFICATE.docx", "RESULT CANCELLATION CERTIFICATE.pdf")
        
        images = convert_from_path("RESULT CANCELLATION CERTIFICATE.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'RESULT CANCELLATION CERTIFICATE'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()
        wind.attributes('-fullscreen',True)
        

        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        # canvas= Canvas(wind, width= 900, height= 900)
        # canvas.pack()
        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")

        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()

        #Load an image in the script
        img= (Image.open("RESULT CANCELLATION CERTIFICATE.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

        




    
        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(), OpenScanDocumentFile()])
        btn_.place(x = 45,
                y=500,
                width=40,
                height=40
                
                        
                        ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(), OpenScanDocumentFile() ])
        btn2_.place(x = 130,
                y=500,
                width=40,
                height=40
                        
                )
        
        wind.mainloop()
        
        # # for Result Cancelation Certificate
    
    if certificate_name_code == '03000313':
        

        R_C_C_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_C_Name, [reg_no_])

        Name_R_C_C = cursor_3.fetchone()[0]

        R_C_C_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_C_fname, [reg_no_])
        

        fname_R_C_C = cursor_3.fetchone()[0]

        R_C_C_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        cursor_4.execute(R_C_C_roll_no, [reg_no_])

        roll_no_R_C_C = cursor_4.fetchone()[0]

        R_C_C_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor_3.execute(R_C_C_reg_no, [fname_R_C_C])

        reg_no_R_C_C = cursor_3.fetchone()[0]
        
        
        
        template_path = os.path.join(main_path, 'RESULT CANCELLATION CERTIFICATE_templ.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        # workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        # worksheet = workbook["Input"]

        to_fill_in = {'Candidate_name' : None,
                    'Father_name': None,
                    'Roll_no' : None,
                    'Registration_no' : None,
                    'Group':None,
                    
                    
                    }

        to_fill_in['Candidate_name'] = Name_R_C_C
        to_fill_in['Dated'] = Dated
        to_fill_in['Father_name'] = fname_R_C_C
        to_fill_in['Registration_no'] = reg_no_R_C_C
        to_fill_in['Roll_no'] = roll_no_R_C_C
        
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'RESULT CANCELLATION CERTIFICATE.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Done with RESULT CANCELLATION CERTIFICATE.docx")
        

        
        # Converting docx present in the same folder
        # as the python file
        # convert("RESULT CANCELLATION CERTIFICATE.docx")
        
        # Converting docx specifying both the input
        # and output paths
        convert("RESULT CANCELLATION CERTIFICATE.docx", "RESULT CANCELLATION CERTIFICATE.pdf")
        
        images = convert_from_path("RESULT CANCELLATION CERTIFICATE.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'RESULT CANCELLATION CERTIFICATE'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        
        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()

    
        #Load an image in the script
        img= (Image.open("RESULT CANCELLATION CERTIFICATE.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)
        
    


        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(), OpenScanDocumentFile() ])
        btn_.place(x = 45,
                    y=500,
                    width=40,
                    height=40
                    
                            
                            ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile()])
        btn2_.place(x = 130,
                    y=500,
                    width=40,
                    height=40
                    
                    )
        
        wind.mainloop()
        
    try:
        computer_tem = '''Select sub_code from ZLedger INNER JOIN ZLedgerSub  on
        ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
        and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ?  and sub_code ='041' '''
        cursor_2.execute(computer_tem , [reg_no_])
        computer_tem = cursor_2.fetchone()[0]
    except:
        computer_tem = 0
        
    print(computer_tem)
                
            # try:
            #     computer_tem = '''Select sub_code from ZLedger INNER JOIN ZLedgerSub  on
            #     ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
            #     and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ?  and sub_code ='041' '''
            #     cursor_2.execute(computer_tem , [reg_no_])
            #     computer_tem = cursor_2.fetchone()[0]
            
                
    # if computer_tem == '041':
            
        # template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION SSC_temp.docx')
        # workbook_path = os.path.join(main_path, 'Template_data.xlsx')
            
    # if computer_tem == '075':
    if certificate_name_code == '02000216' and computer_tem == '075':
        
        
            # print("i am in")
            
        
            # # For Application form for migration request

    # .....
            R_C_A_X_roll_no = '''select roll_no from ZLedger where reg_no = ?
    '''
            cursor_2.execute(R_C_A_X_roll_no, [reg_no_])

            roll_no_R_C_A_X = cursor_2.fetchone()[0]
            # print(roll_no_R_C_A_X)
            
            R_C_A_X_group = '''Select grp_code from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  '''
            cursor_2.execute(R_C_A_X_group, [roll_no_R_C_A_X])

            R_C_A_X_group = cursor_2.fetchone()[0]
            if R_C_A_X_group == '1':
                R_C_A_X_group = "SCIENCE"
            elif R_C_A_X_group == '2':
                R_C_A_X_group = "HUMANITIES"
            elif R_C_A_X_group =='3':
                R_C_A_X_group ="TECHNICAL"
            elif R_C_A_X_group == '4':
                R_C_A_X_group = "MATRIC TECH"
                

            R_C_A_X_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
            cursor.execute(R_C_A_X_Name, [reg_no_])

            Name_R_C_A_X = cursor.fetchone()[0]
            print(Name_R_C_A_X)  



            R_C_A_X_fname = '''Select fname from ZReg WHERE reg_no = ?'''
            cursor.execute(R_C_A_X_fname, [reg_no_])

            fname_R_C_A_X = cursor.fetchone()[0]

            R_C_A_X_reg_no = '''Select reg_no from ZReg where fname = ?'''
            cursor.execute(R_C_A_X_reg_no, [fname_R_C_A_X])

            reg_no_R_C_A_X = cursor.fetchone()[0]



            R_C_A_X_Year = '''Select year from ZReg WHERE reg_no = ?'''
            cursor.execute(R_C_A_X_Year, [reg_no_])

            Year_R_C_A_X = cursor.fetchone()[0]

            R_C_A_X_Institution = '''select inst_desc from ZReg where reg_no = ?'''
            cursor.execute(R_C_A_X_Institution, [reg_no_])
            Institution_R_C_A_X  = cursor.fetchone()[0]


            # total_marks = '''Select marks_obt from ZLedgerII where reg_no = ?  ''' 
            # cursor_2.execute(total_marks, [reg_no_])
            # Institution_M_C = cursor_2.fetchone()[0]

            # R_C_A_X_sts = '''Select pass_fail_status from ZLedgerII where reg = ?'''
            # cursor_2.execute(R_C_A_X_sts, [reg_no_])
            # R_C_A_X_sts= cursor_2.fetchone()[0]

            R_C_A_X_eng = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '001'
    '''
            cursor_2.execute(R_C_A_X_eng, [reg_no_])
            R_C_A_X_eng= cursor_2.fetchone()[0]
            print(R_C_A_X_eng)


            R_C_A_X_urdu = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '002' '''
            cursor_2.execute(R_C_A_X_urdu, [reg_no_])
            R_C_A_X_urdu= cursor_2.fetchone()[0]
            try:
                R_C_A_X_isl = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '108' '''
                cursor_2.execute(R_C_A_X_isl, [reg_no_])
                R_C_A_X_isl= cursor_2.fetchone()[0]
            except:
                R_C_A_X_isl = 0
                

            try:
                R_C_A_X_pakS = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '106' '''
                cursor_2.execute(R_C_A_X_pakS, [reg_no_])
                R_C_A_X_pakS= cursor_2.fetchone()[0]
            except:
                R_C_A_X_pakS = 0
            try:
                R_C_A_X_math = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '010' '''
                cursor_2.execute(R_C_A_X_math, [reg_no_])
                R_C_A_X_math= cursor_2.fetchone()[0]
            except:
                R_C_A_X_math = 0
            print(R_C_A_X_math)

            R_C_A_X_phy = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy, [reg_no_])
            R_C_A_X_phy= cursor_2.fetchone()[0]
            
            R_C_A_X_phy_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy_p, [roll_no_R_C_A_X ])
            R_C_A_X_phy_p= cursor_2.fetchone()[0]

            R_C_A_X_che = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '012'  '''
            cursor_2.execute(R_C_A_X_che, [reg_no_])
            R_C_A_X_che= cursor_2.fetchone()[0]
            
            
            R_C_A_X_che_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
            cursor_2.execute(R_C_A_X_che_p, [roll_no_R_C_A_X ])
            R_C_A_X_che_p= cursor_2.fetchone()[0]
            try:
                R_C_A_X_bio = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code ='041' '''
                cursor_2.execute(R_C_A_X_bio, [reg_no_])
                R_C_A_X_bio= cursor_2.fetchone()[0]
            except:
                R_C_A_X_bio = 0
                
            try:
                R_C_A_X_bio_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
                cursor_2.execute(R_C_A_X_bio_p, [roll_no_R_C_A_X ])
                R_C_A_X_bio_p= cursor_2.fetchone()[0]
            except:
                R_C_A_X_bio_p = 0

            try:
                R_C_A_X_obt = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '001' '''
                cursor_2.execute(R_C_A_X_obt, [reg_no_])
                R_C_A_X_obt= cursor_2.fetchone()[0]
            except:
                R_C_A_X_obt = 0
                
            try:
                R_C_A_X_obt_2 = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '013' and app_sts ='2' and  ZLedger.reg_no = ?'''
                cursor_2.execute(R_C_A_X_obt_2, [reg_no_])
                R_C_A_X_obt_2= cursor_2.fetchone()[0]
            except:
                R_C_A_X_obt_2 = 0
                
                
            R_C_A_X_eng2 = ''' Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '001' '''
            cursor_2.execute(R_C_A_X_eng2, [roll_no_R_C_A_X ])
            R_C_A_X_eng2= cursor_2.fetchone()[0]
            print(R_C_A_X_eng2)


            R_C_A_X_urdu2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '002' '''
            cursor_2.execute(R_C_A_X_urdu2, [roll_no_R_C_A_X ])
            R_C_A_X_urdu2= cursor_2.fetchone()[0] 
            try:
                R_C_A_X_isl2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '008' '''
                cursor_2.execute(R_C_A_X_isl2, [roll_no_R_C_A_X ])
                R_C_A_X_isl2= cursor_2.fetchone()[0]
            except:
                R_C_A_X_isl2 = 0
                

            R_C_A_X_pakS2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '106' '''
            cursor_2.execute(R_C_A_X_pakS2, [roll_no_R_C_A_X ])
            R_C_A_X_pakS2= cursor_2.fetchone()[0]

            try:
                R_C_A_X_math2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '010' '''
                cursor_2.execute(R_C_A_X_math2, [roll_no_R_C_A_X ])
                R_C_A_X_math2= cursor_2.fetchone()[0]
            except:
                R_C_A_X_math2 = 0
            print(R_C_A_X_math2)

            R_C_A_X_phy2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy2, [roll_no_R_C_A_X ])
            R_C_A_X_phy2= cursor_2.fetchone()[0]
            
            try:
                R_C_A_X_phy2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
                cursor_2.execute(R_C_A_X_phy2_p, [roll_no_R_C_A_X ])
                R_C_A_X_phy2_p= cursor_2.fetchone()[0]
                
            except:
                R_C_A_X_phy2_p = 0

            try:
                R_C_A_X_che2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
                cursor_2.execute(R_C_A_X_che2, [roll_no_R_C_A_X ])
                R_C_A_X_che2= cursor_2.fetchone()[0]
            
            except:
                R_C_A_X_che2_p = 0

            try:
                R_C_A_X_che2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
                cursor_2.execute(R_C_A_X_che2_p, [roll_no_R_C_A_X ])
                R_C_A_X_che2_p= cursor_2.fetchone()[0]
            
            except:
                R_C_A_X_che2_p = 0
                
            try:
                R_C_A_X_bio2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
                cursor_2.execute(R_C_A_X_bio2, [roll_no_R_C_A_X ])
                R_C_A_X_bio2= cursor_2.fetchone()[0]
            except:
                R_C_A_X_bio2_p = 0
            try:
                R_C_A_X_bio2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
                cursor_2.execute(R_C_A_X_bio2_p, [roll_no_R_C_A_X ])
                R_C_A_X_bio2_p= cursor_2.fetchone()[0]
            except:
                R_C_A_X_bio2_p = 0
                
            
                

                
            
            

            # R_C_A_X_total = '''SELECT total FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
            # cursor_2.execute(R_C_A_X_total, [reg_no_])
            R_C_A_X_total= 1100
            
            # R_C_A_X_idf = '''SELECT id_mark FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
            # cursor_2.execute(R_C_A_X_idf, [reg_no_])
            R_C_A_X_idf="Nothing"
            
            R_C_A_X_sts_e = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_sts_e, [reg_no_])
            R_C_A_X_sts_e= cursor_2.fetchone()[0]
            if(R_C_A_X_sts_e== '1'):
                R_C_A_X_sts_e = "Pass"
            
            else:
                
                R_C_A_X_sts_e = "Fail"
                
            R_C_A_X_sts_u = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '002' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_sts_u, [reg_no_])
            R_C_A_X_sts_u= cursor_2.fetchone()[0]
            if(R_C_A_X_sts_u== '1'):
                R_C_A_X_sts_u = "Pass"
            
            else:
                
                R_C_A_X_sts_u = "Fail"
                
            R_C_A_X_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
            cursor.execute(R_C_A_X_DOB, [reg_no_])

            DOB_R_C_A_X = cursor.fetchone()[0]
            
            
            R_C_A_X_sts_i = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '108' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_sts_i, [reg_no_])
            R_C_A_X_sts_i= cursor_2.fetchone()[0]
            if(R_C_A_X_sts_i== '1'):
                R_C_A_X_sts_i = "Pass"
            
            else:
                
                R_C_A_X_sts_i = "Fail"
                
            R_C_A_X_sts_pk = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
    ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
    and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '106' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_sts_pk, [reg_no_])
            R_C_A_X_sts_pk= cursor_2.fetchone()[0]
            if(R_C_A_X_sts_pk== '1'):
                R_C_A_X_sts_pk = "Pass"
            
            else:
                
                R_C_A_X_sts_pk = "Fail"
                
                
            


            # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
            # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
            # R_C_A_X_eng= cursor_2.fetchone()[0]

            # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
            # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
            # R_C_A_X_eng= cursor_2.fetchone()[0]
            









            
            template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION SSC_temp.docx')
            workbook_path = os.path.join(main_path, 'Template_data.xlsx')
                

            workbook = load_workbook(workbook_path)
            template = DocxTemplate(template_path)
            worksheet = workbook["Input"]

            to_fill_in = {
                        'Roll_no' : None,
                        'Group' : None,
                        'Registration_no' : None,
                        'Certificate_no' : None,
                        'Candidate_name':None,
                        'Father_name': None,
                        'Institution': None,
                        'sts_e':None,
                        'sts_u':None,
                        'sts_i':None,
                        'sts_pk':None,
                        'sts_m':None,
                        'sts_ph':None,
                        'sts_ch':None,
                        'sts_bio':None,
                        'sts_ph_p':None,
                        'sts_ch_p':None,
                        'sts_bio_p':None,
                        
                        # 'R_C_A_X_eng2':None,
                        # 'R_C_A_X_urdu2':None,
                        # 'R_C_A_X_isl2':None,
                        # 'R_C_A_X_paskS2':None,
                        # 'R_C_A_X_math2':None,
                        # 'R_C_A_X_phy2':None,
                        # 'R_C_A_X_che2':None,
                        # 'R_C_A_X_bio2':None,
                        
                        # 'R_C_A_X_phy2_p':None,
                        # 'R_C_A_X_che2_p':None,
                        # 'R_C_A_X_bio2_p':None,
                        
                        'R_C_A_X_eng':None,
                        'R_C_A_X_urdu':None,
                        'R_C_A_X_isl':None,
                        'R_C_A_X_paskS':None,
                        'R_C_A_X_math':None,
                        'R_C_A_X_phy':None,
                        'R_C_A_X_che':None,
                        'R_C_A_X_bio':None,
                        
                        'R_C_A_X_eng_T':None,
                        'R_C_A_X_urdu_T':None,
                        'R_C_A_X_isl_T':None,
                        'R_C_A_X_paskS_T':None,
                        'R_C_A_X_math_T':None,
                        'R_C_A_X_phy_T':None,
                        'R_C_A_X_che_T':None,
                        'R_C_A_X_bio_T':None,
                        
                        'R_C_A_X_phy_p':None,
                        'R_C_A_X_che_p':None,
                        'R_C_A_X_bio_p':None,
                        
                        'total_marks':None,
                        'idf':None,
                        
                        
                        
                        'total_mark_obt':None,
                        'total_mark_obt_2':None,
                        
                        
                        
                        
                        
                        'Dated' : None,
                        'DOB': None,
                        'Examination': None,
                        'Session': None,
                        'year': None,
                        'Status': None,
                        
                        'year':None,
                        'Year':None,
                        
                        
                        }




            to_fill_in['Roll_no'] = roll_no_R_C_A_X
            to_fill_in['Group'] = R_C_A_X_group
            to_fill_in['Registration_no'] = reg_no_
            to_fill_in['Certificate_no'] = 1223344
            to_fill_in['Candidate_name'] =Name_R_C_A_X
            to_fill_in['Father_name'] = fname_R_C_A_X
            to_fill_in['Institution'] = Institution_R_C_A_X 
            to_fill_in['total_mark_obt_2'] = R_C_A_X_obt_2
            # R_C_A_X_obt = int(R_C_A_X_obt_2)+int(R_C_A_X_obt)

            to_fill_in['total_mark_obt'] = R_C_A_X_obt
            
            
            
            to_fill_in['Dated']=  Dated
            # to_fill_in['year']= year
            to_fill_in['Year'] = year_curr
            
            to_fill_in['sts_e']=R_C_A_X_sts_e
            to_fill_in['sts_u']=R_C_A_X_sts_u
            to_fill_in['sts_i']=R_C_A_X_sts_i
            to_fill_in['sts_pk']=R_C_A_X_sts_pk
            to_fill_in['sts_m']=R_C_A_X_sts_e
            to_fill_in['sts_ph']=R_C_A_X_sts_e
            to_fill_in['sts_ch']=R_C_A_X_sts_e
            to_fill_in['sts_bio']=R_C_A_X_sts_e
            
            
            to_fill_in['R_C_A_X_eng']= R_C_A_X_eng 
            to_fill_in['R_C_A_X_urdu']=R_C_A_X_urdu
            to_fill_in['R_C_A_X_isl']= R_C_A_X_isl
            to_fill_in['R_C_A_X_pakS']= R_C_A_X_pakS
            to_fill_in['R_C_A_X_math']=R_C_A_X_math
            to_fill_in['R_C_A_X_phy']= R_C_A_X_phy
            to_fill_in['R_C_A_X_che']=R_C_A_X_che
            to_fill_in['R_C_A_X_bio']= R_C_A_X_bio
            
            # to_fill_in['R_C_A_X_phy_p']= R_C_A_X_phy_p
            # to_fill_in['R_C_A_X_che_p']=R_C_A_X_che_p
            # to_fill_in['R_C_A_X_bio_p']= R_C_A_X_bio_p
            
            # to_fill_in['R_C_A_X_eng2']= R_C_A_X_eng2
            # to_fill_in['R_C_A_X_urdu2']=R_C_A_X_urdu2
            # to_fill_in['R_C_A_X_isl2']= R_C_A_X_isl2
            # to_fill_in['R_C_A_X_pakS2']= R_C_A_X_pakS2
            # to_fill_in['R_C_A_X_math2']=R_C_A_X_math2
            # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
            # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
            # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
            
            
            #  addition for the total Values
            
            # total_eng = int(R_C_A_X_eng2)+int(R_C_A_X_eng)
            # total_urdu =  int(R_C_A_X_urdu2)+int(R_C_A_X_urdu)
            # total_isl = int(R_C_A_X_isl2)+int(R_C_A_X_isl)
            # total_paks = int(R_C_A_X_pakS2)+int(R_C_A_X_pakS)
            # total_math = int(R_C_A_X_math2)+int(R_C_A_X_math)
            # total_phy = int(R_C_A_X_phy2)+int(R_C_A_X_phy)
            # total_che =int(R_C_A_X_che2)+int(R_C_A_X_che)
            # total_bio = int(R_C_A_X_bio2)+int(R_C_A_X_bio)
            
            # # end
            # to_fill_in['R_C_A_X_eng_T']= total_eng 
            # to_fill_in['R_C_A_X_urdu_T']=total_urdu
            # to_fill_in['R_C_A_X_isl_T']= total_isl
            # to_fill_in['R_C_A_X_pakS_T']= total_paks
            # to_fill_in['R_C_A_X_math_T']=total_math
            # to_fill_in['R_C_A_X_phy_T']=  total_phy
            # to_fill_in['R_C_A_X_che_T']=total_che
            # to_fill_in['R_C_A_X_bio_T']= total_bio
            
            # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
            # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
            # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
            
            # to_fill_in['R_C_A_X_phy2_p']= R_C_A_X_phy2_p
            # to_fill_in['R_C_A_X_che2_p']=R_C_A_X_che2_p
            # to_fill_in['R_C_A_X_bio2_p']= R_C_A_X_bio2_p
            
            
            to_fill_in['total_marks']= 510


            to_fill_in['Session'] = "Final Session"
            to_fill_in['year'] =  2002
            to_fill_in['Status'] = "pass"
            to_fill_in['Institution'] = Institution_R_C_A_X
            to_fill_in['DOB'] = DOB_R_C_A_X 
            to_fill_in['idf']= R_C_A_X_idf
                
                
            # Fill in all the keys defined in the word document using the dictionary.
            # The keys in de word document are identified by the {{}}symbols.
            template.render(to_fill_in)
            # Output the file to a docx document.
            filename = 'Result Card ANNUAL EXAMINATION SSC.docx'
            filled_path = os.path.join(main_path, filename)
            template.save(filled_path)
            print("Result Card ANNUAL EXAMINATION SSC.docx")
            doc = Document('Result Card ANNUAL EXAMINATION SSC.docx')
            tables = doc.tables
            p = tables[0].rows[0].cells[0].add_paragraph()
            # r = p.add_run()
            # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
            # p = tables[0].rows[0].cells[0].add_paragraph()
            r = p.add_run()
            data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
            cursor_2.execute(data, [reg_no_])
            data = cursor_2.fetchone()[0]
            r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
            # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
            doc.save('Result Card ANNUAL EXAMINATION SSC.docx')

            convert("Result Card ANNUAL EXAMINATION SSC.docx", "Result Card ANNUAL EXAMINATION SSC.pdf")

            images = convert_from_path("Result Card ANNUAL EXAMINATION SSC.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
            for i, image in enumerate(images):
                fname = 'Result Card ANNUAL EXAMINATION SSC'+'.png'
                image.save(fname, "PNG")
                
            # Python3 program to convert docx to 
                

            #Import the required Libraries

            #Create an instance of tkinter frame
            wind = tk.Toplevel()

            #Set the geometry of tkinter frame
            wind.geometry("1600x1600")
            wind.attributes('-fullscreen',True)
            IMAGE_PATH = 'sir-ok.jpg'
            # WIDTH, HEIGTH = 600, 600

            # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

            img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
            lbl = tk.Label(wind, image=img)
            lbl.img = img  # Keep a reference in case this code put is in a function.
            lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        

            #Create a canvas
            canvas= Canvas(wind, width= 900, height= 900)
            canvas.pack()
            
            # Add image file
            # bg = PhotoImage(file = "sir-ok.jpg")
            
            # # Show image using label
            # label1 = Label( root, image = bg)
            # label1.place(x = 0, y = 0)
            
            
            
    
            # # frame = Frame(wind, width=600, height=400)
            # frame.pack()
            # frame.place(anchor='center', relx=0.5, rely=0.5)

            # # Create an object of tkinter ImageTk
            # img = ImageTk.PhotoImage(Image.open("sir-ok.jpg"))

            # # Create a Label Widget to display the text or Image
            # label = Label(frame, image = img)
            # label.pack()


            #Load an image in the script
            img= (Image.open("Result Card ANNUAL EXAMINATION SSC.png"))

            #Resize the Image using resize method
            resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
            new_image= ImageTk.PhotoImage(resized_image)

            #Add image to the Canvas Items
            canvas.create_image(10,10, anchor=NW, image=new_image)

            


            btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),  OpenScanDocumentFile()  ])
            btn_.place(x = 45,
                    y=500,
                    width=40,
                    height=40
                    
                            
                            ) 

            btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile() ])
            btn2_.place(x = 130,
                    y=500,
                    width=40,
                    height=40
                    
                    )
            
            
            
            

            wind.mainloop()
            
            # After second testing
            #  and computer_tem == '041'
    
    if certificate_name_code == '02000216':
        
        
        
            # print("i am in")
            
        
            # # For Application form for migration request

    # .....
        R_C_A_X_roll_no = '''select roll_no from ZLedger where reg_no = ?
'''
        cursor_2.execute(R_C_A_X_roll_no, [reg_no_])

        roll_no_R_C_A_X = cursor_2.fetchone()[0]
        # print(roll_no_R_C_A_X)
        
        R_C_A_X_group = '''Select grp_code from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  '''
        cursor_2.execute(R_C_A_X_group, [roll_no_R_C_A_X])

        R_C_A_X_group = cursor_2.fetchone()[0]
        if R_C_A_X_group == '1':
            R_C_A_X_group = "SCIENCE"
        elif R_C_A_X_group == '2':
            R_C_A_X_group = "HUMANITIES"
        elif R_C_A_X_group =='3':
            R_C_A_X_group ="TECHNICAL"
        elif R_C_A_X_group == '4':
            R_C_A_X_group = "MATRIC TECH"
            

        R_C_A_X_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_Name, [reg_no_])

        Name_R_C_A_X = cursor.fetchone()[0]
        print(Name_R_C_A_X)  



        R_C_A_X_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_fname, [reg_no_])

        fname_R_C_A_X = cursor.fetchone()[0]

        R_C_A_X_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor.execute(R_C_A_X_reg_no, [fname_R_C_A_X])

        reg_no_R_C_A_X = cursor.fetchone()[0]



        R_C_A_X_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_Year, [reg_no_])

        Year_R_C_A_X = cursor.fetchone()[0]

        R_C_A_X_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor.execute(R_C_A_X_Institution, [reg_no_])
        Institution_R_C_A_X  = cursor.fetchone()[0]


        # total_marks = '''Select marks_obt from ZLedgerII where reg_no = ?  ''' 
        # cursor_2.execute(total_marks, [reg_no_])
        # Institution_M_C = cursor_2.fetchone()[0]

        # R_C_A_X_sts = '''Select pass_fail_status from ZLedgerII where reg = ?'''
        # cursor_2.execute(R_C_A_X_sts, [reg_no_])
        # R_C_A_X_sts= cursor_2.fetchone()[0]

        R_C_A_X_eng = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '001'
'''
        cursor_2.execute(R_C_A_X_eng, [reg_no_])
        R_C_A_X_eng= cursor_2.fetchone()[0]
        print(R_C_A_X_eng)


        R_C_A_X_urdu = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '002' '''
        cursor_2.execute(R_C_A_X_urdu, [reg_no_])
        R_C_A_X_urdu= cursor_2.fetchone()[0]
        try:
            R_C_A_X_isl = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '108' '''
            cursor_2.execute(R_C_A_X_isl, [reg_no_])
            R_C_A_X_isl= cursor_2.fetchone()[0]
        except:
            R_C_A_X_isl = 0
            

        try:
            R_C_A_X_pakS = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '106' '''
            cursor_2.execute(R_C_A_X_pakS, [reg_no_])
            R_C_A_X_pakS= cursor_2.fetchone()[0]
        except:
            R_C_A_X_pakS = 0
        try:
            R_C_A_X_math = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '010' '''
            cursor_2.execute(R_C_A_X_math, [reg_no_])
            R_C_A_X_math= cursor_2.fetchone()[0]
        except:
            R_C_A_X_math = 0
        print(R_C_A_X_math)

        R_C_A_X_phy = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '011' '''
        cursor_2.execute(R_C_A_X_phy, [reg_no_])
        R_C_A_X_phy= cursor_2.fetchone()[0]
        
        R_C_A_X_phy_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
        cursor_2.execute(R_C_A_X_phy_p, [roll_no_R_C_A_X ])
        R_C_A_X_phy_p= cursor_2.fetchone()[0]

        R_C_A_X_che = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '012'  '''
        cursor_2.execute(R_C_A_X_che, [reg_no_])
        R_C_A_X_che= cursor_2.fetchone()[0]
        
        
        R_C_A_X_che_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
        cursor_2.execute(R_C_A_X_che_p, [roll_no_R_C_A_X ])
        R_C_A_X_che_p= cursor_2.fetchone()[0]
        try:
            R_C_A_X_bio = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code ='041' '''
            cursor_2.execute(R_C_A_X_bio, [reg_no_])
            R_C_A_X_bio= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio = 0
            
        try:
            R_C_A_X_bio_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio_p= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio_p = 0

        try:
            R_C_A_X_obt = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.reg_no = ? and sub_code = '001' '''
            cursor_2.execute(R_C_A_X_obt, [reg_no_])
            R_C_A_X_obt= cursor_2.fetchone()[0]
        except:
            R_C_A_X_obt = 0
            
        try:
            R_C_A_X_obt_2 = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '013' and app_sts ='2' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_obt_2, [reg_no_])
            R_C_A_X_obt_2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_obt_2 = 0
            
            
        R_C_A_X_eng2 = ''' Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '001' '''
        cursor_2.execute(R_C_A_X_eng2, [roll_no_R_C_A_X ])
        R_C_A_X_eng2= cursor_2.fetchone()[0]
        print(R_C_A_X_eng2)


        R_C_A_X_urdu2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '002' '''
        cursor_2.execute(R_C_A_X_urdu2, [roll_no_R_C_A_X ])
        R_C_A_X_urdu2= cursor_2.fetchone()[0] 
        try:
            R_C_A_X_isl2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '008' '''
            cursor_2.execute(R_C_A_X_isl2, [roll_no_R_C_A_X ])
            R_C_A_X_isl2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_isl2 = 0
            

        R_C_A_X_pakS2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '106' '''
        cursor_2.execute(R_C_A_X_pakS2, [roll_no_R_C_A_X ])
        R_C_A_X_pakS2= cursor_2.fetchone()[0]

        try:
            R_C_A_X_math2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '010' '''
            cursor_2.execute(R_C_A_X_math2, [roll_no_R_C_A_X ])
            R_C_A_X_math2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_math2 = 0
        print(R_C_A_X_math2)

        R_C_A_X_phy2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
        cursor_2.execute(R_C_A_X_phy2, [roll_no_R_C_A_X ])
        R_C_A_X_phy2= cursor_2.fetchone()[0]
        
        try:
            R_C_A_X_phy2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy2_p, [roll_no_R_C_A_X ])
            R_C_A_X_phy2_p= cursor_2.fetchone()[0]
            
        except:
            R_C_A_X_phy2_p = 0

        try:
            R_C_A_X_che2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_2.execute(R_C_A_X_che2, [roll_no_R_C_A_X ])
            R_C_A_X_che2= cursor_2.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0

        try:
            R_C_A_X_che2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_2.execute(R_C_A_X_che2_p, [roll_no_R_C_A_X ])
            R_C_A_X_che2_p= cursor_2.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0
            
        try:
            R_C_A_X_bio2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio2, [roll_no_R_C_A_X ])
            R_C_A_X_bio2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
        try:
            R_C_A_X_bio2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio2_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio2_p= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
            
        
            

            
        
        

        # R_C_A_X_total = '''SELECT total FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_total, [reg_no_])
        R_C_A_X_total= 1100
        
        # R_C_A_X_idf = '''SELECT id_mark FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_idf, [reg_no_])
        R_C_A_X_idf="Nothing"
        
        R_C_A_X_sts_e = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_e, [reg_no_])
        R_C_A_X_sts_e= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_e== '1'):
            R_C_A_X_sts_e = "Pass"
        
        else:
            
            R_C_A_X_sts_e = "Fail"
            
        R_C_A_X_sts_u = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '002' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_u, [reg_no_])
        R_C_A_X_sts_u= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_u== '1'):
            R_C_A_X_sts_u = "Pass"
        
        else:
            
            R_C_A_X_sts_u = "Fail"
            
        R_C_A_X_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_DOB, [reg_no_])

        DOB_R_C_A_X = cursor.fetchone()[0]
        
        
        R_C_A_X_sts_i = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '108' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_i, [reg_no_])
        R_C_A_X_sts_i= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_i== '1'):
            R_C_A_X_sts_i = "Pass"
        
        else:
            
            R_C_A_X_sts_i = "Fail"
            
        R_C_A_X_sts_pk = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '106' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_pk, [reg_no_])
        R_C_A_X_sts_pk= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_pk== '1'):
            R_C_A_X_sts_pk = "Pass"
        
        else:
            
            R_C_A_X_sts_pk = "Fail"
            
            
        


        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]

        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]
        









        
        template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION SSC_temp.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')
            

        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {
                    'Roll_no' : None,
                    'Group' : None,
                    'Registration_no' : None,
                    'Certificate_no' : None,
                    'Candidate_name':None,
                    'Father_name': None,
                    'Institution': None,
                    'sts_e':None,
                    'sts_u':None,
                    'sts_i':None,
                    'sts_pk':None,
                    'sts_m':None,
                    'sts_ph':None,
                    'sts_ch':None,
                    'sts_bio':None,
                    'sts_ph_p':None,
                    'sts_ch_p':None,
                    'sts_bio_p':None,
                    
                    # 'R_C_A_X_eng2':None,
                    # 'R_C_A_X_urdu2':None,
                    # 'R_C_A_X_isl2':None,
                    # 'R_C_A_X_paskS2':None,
                    # 'R_C_A_X_math2':None,
                    # 'R_C_A_X_phy2':None,
                    # 'R_C_A_X_che2':None,
                    # 'R_C_A_X_bio2':None,
                    
                    # 'R_C_A_X_phy2_p':None,
                    # 'R_C_A_X_che2_p':None,
                    # 'R_C_A_X_bio2_p':None,
                    
                    'R_C_A_X_eng':None,
                    'R_C_A_X_urdu':None,
                    'R_C_A_X_isl':None,
                    'R_C_A_X_paskS':None,
                    'R_C_A_X_math':None,
                    'R_C_A_X_phy':None,
                    'R_C_A_X_che':None,
                    'R_C_A_X_bio':None,
                    
                    'R_C_A_X_eng_T':None,
                    'R_C_A_X_urdu_T':None,
                    'R_C_A_X_isl_T':None,
                    'R_C_A_X_paskS_T':None,
                    'R_C_A_X_math_T':None,
                    'R_C_A_X_phy_T':None,
                    'R_C_A_X_che_T':None,
                    'R_C_A_X_bio_T':None,
                    
                    'R_C_A_X_phy_p':None,
                    'R_C_A_X_che_p':None,
                    'R_C_A_X_bio_p':None,
                    
                    'total_marks':None,
                    'idf':None,
                    
                    
                    
                    'total_mark_obt':None,
                    'total_mark_obt_2':None,
                    
                    
                    
                    
                    
                    'Dated' : None,
                    'DOB': None,
                    'Examination': None,
                    'Session': None,
                    'year': None,
                    'Status': None,
                    
                    'year':None,
                    'Year':None,
                    
                    
                    }




        to_fill_in['Roll_no'] = roll_no_R_C_A_X
        to_fill_in['Group'] = R_C_A_X_group
        to_fill_in['Registration_no'] = reg_no_
        to_fill_in['Certificate_no'] = 1223344
        to_fill_in['Candidate_name'] =Name_R_C_A_X
        to_fill_in['Father_name'] = fname_R_C_A_X
        to_fill_in['Institution'] = Institution_R_C_A_X 
        to_fill_in['total_mark_obt_2'] = R_C_A_X_obt_2
        # R_C_A_X_obt = int(R_C_A_X_obt_2)+int(R_C_A_X_obt)

        to_fill_in['total_mark_obt'] = R_C_A_X_obt
        
        
        
        to_fill_in['Dated']=  Dated
        # to_fill_in['year']= year
        to_fill_in['Year'] = year_curr
        
        to_fill_in['sts_e']=R_C_A_X_sts_e
        to_fill_in['sts_u']=R_C_A_X_sts_u
        to_fill_in['sts_i']=R_C_A_X_sts_i
        to_fill_in['sts_pk']=R_C_A_X_sts_pk
        to_fill_in['sts_m']=R_C_A_X_sts_e
        to_fill_in['sts_ph']=R_C_A_X_sts_e
        to_fill_in['sts_ch']=R_C_A_X_sts_e
        to_fill_in['sts_bio']=R_C_A_X_sts_e
        
        
        to_fill_in['R_C_A_X_eng']= R_C_A_X_eng 
        to_fill_in['R_C_A_X_urdu']=R_C_A_X_urdu
        to_fill_in['R_C_A_X_isl']= R_C_A_X_isl
        to_fill_in['R_C_A_X_pakS']= R_C_A_X_pakS
        to_fill_in['R_C_A_X_math']=R_C_A_X_math
        to_fill_in['R_C_A_X_phy']= R_C_A_X_phy
        to_fill_in['R_C_A_X_che']=R_C_A_X_che
        to_fill_in['R_C_A_X_bio']= R_C_A_X_bio
        
        # to_fill_in['R_C_A_X_phy_p']= R_C_A_X_phy_p
        # to_fill_in['R_C_A_X_che_p']=R_C_A_X_che_p
        # to_fill_in['R_C_A_X_bio_p']= R_C_A_X_bio_p
        
        # to_fill_in['R_C_A_X_eng2']= R_C_A_X_eng2
        # to_fill_in['R_C_A_X_urdu2']=R_C_A_X_urdu2
        # to_fill_in['R_C_A_X_isl2']= R_C_A_X_isl2
        # to_fill_in['R_C_A_X_pakS2']= R_C_A_X_pakS2
        # to_fill_in['R_C_A_X_math2']=R_C_A_X_math2
        # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        
        #  addition for the total Values
        
        # total_eng = int(R_C_A_X_eng2)+int(R_C_A_X_eng)
        # total_urdu =  int(R_C_A_X_urdu2)+int(R_C_A_X_urdu)
        # total_isl = int(R_C_A_X_isl2)+int(R_C_A_X_isl)
        # total_paks = int(R_C_A_X_pakS2)+int(R_C_A_X_pakS)
        # total_math = int(R_C_A_X_math2)+int(R_C_A_X_math)
        # total_phy = int(R_C_A_X_phy2)+int(R_C_A_X_phy)
        # total_che =int(R_C_A_X_che2)+int(R_C_A_X_che)
        # total_bio = int(R_C_A_X_bio2)+int(R_C_A_X_bio)
        
        # # end
        # to_fill_in['R_C_A_X_eng_T']= total_eng 
        # to_fill_in['R_C_A_X_urdu_T']=total_urdu
        # to_fill_in['R_C_A_X_isl_T']= total_isl
        # to_fill_in['R_C_A_X_pakS_T']= total_paks
        # to_fill_in['R_C_A_X_math_T']=total_math
        # to_fill_in['R_C_A_X_phy_T']=  total_phy
        # to_fill_in['R_C_A_X_che_T']=total_che
        # to_fill_in['R_C_A_X_bio_T']= total_bio
        
        # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        # to_fill_in['R_C_A_X_phy2_p']= R_C_A_X_phy2_p
        # to_fill_in['R_C_A_X_che2_p']=R_C_A_X_che2_p
        # to_fill_in['R_C_A_X_bio2_p']= R_C_A_X_bio2_p
        
        
        to_fill_in['total_marks']= 510


        to_fill_in['Session'] = "Final Session"
        to_fill_in['year'] =  2002
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_R_C_A_X
        to_fill_in['DOB'] = DOB_R_C_A_X 
        to_fill_in['idf']= R_C_A_X_idf
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'Result Card ANNUAL EXAMINATION SSC.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Result Card ANNUAL EXAMINATION SSC.docx")
        doc = Document('Result Card ANNUAL EXAMINATION SSC.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_2.execute(data, [reg_no_])
        data = cursor_2.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('Result Card ANNUAL EXAMINATION SSC.docx')

        convert("Result Card ANNUAL EXAMINATION SSC.docx", "Result Card ANNUAL EXAMINATION SSC.pdf")

        images = convert_from_path("Result Card ANNUAL EXAMINATION SSC.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'Result Card ANNUAL EXAMINATION SSC'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
    

        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        
        # Add image file
        # bg = PhotoImage(file = "sir-ok.jpg")
        
        # # Show image using label
        # label1 = Label( root, image = bg)
        # label1.place(x = 0, y = 0)
        
        
        

        # # frame = Frame(wind, width=600, height=400)
        # frame.pack()
        # frame.place(anchor='center', relx=0.5, rely=0.5)

        # # Create an object of tkinter ImageTk
        # img = ImageTk.PhotoImage(Image.open("sir-ok.jpg"))

        # # Create a Label Widget to display the text or Image
        # label = Label(frame, image = img)
        # label.pack()


        #Load an image in the script
        img= (Image.open("Result Card ANNUAL EXAMINATION SSC.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

        


        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),  OpenScanDocumentFile()  ])
        btn_.place(x = 45,
                y=500,
                width=40,
                height=40
                
                        
                        ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile() ])
        btn2_.place(x = 130,
                y=500,
                width=40,
                height=40
                
                )
        
        
        
        

        wind.mainloop()
        
            
            
            
            
            
        

    
    #Condition for selecting the template file which print 
    # if(certificate == "MIGRATION CERTIFICATE"):
        
    # template_path = os.path.join(main_path, 'MIGRATION CERTIFIACTE_templ.docx')
    # workbook_path = os.path.join(main_path, 'Template_data.xlsx')

    # workbook = load_workbook(workbook_path)
    # template = DocxTemplate(template_path)
    # worksheet = workbook["Input"]

    # to_fill_in = {'Candidate_name' : None,
    #             'Dated' : None,
    #             'DOB': None,
    #             'Father_name': None,
    #             'Registration_no' : None,
    #             'Roll_no' : None,
    #             'Examination': None,
    #             'Session': None,
    #             'year': None,
    #             'Status': None,
    #             'Institution': None
                
    #             }




    # to_fill_in['Candidate_name'] = Name_M_C
    # to_fill_in['DOB'] = DOB_M_C
    # to_fill_in['Dated'] = Dated
    # to_fill_in['Father_name'] = fname_M_C
    # to_fill_in['Registration_no'] = reg_no_M_C
    # to_fill_in['Roll_no'] = roll_no_M_C
    # to_fill_in['Examination'] = "Fedral board"
    # to_fill_in['Session'] = "Final Session"
    # to_fill_in['year'] =  Year_M_C
    # to_fill_in['Status'] = "pass"
    # to_fill_in['Institution'] = Institution_M_C
        
        
    # # Fill in all the keys defined in the word document using the dictionary.
    # # The keys in de word document are identified by the {{}}symbols.
    # template.render(to_fill_in)
    # # Output the file to a docx document.
    # filename = 'MIGRATION CERTIFIACTE.docx'
    # filled_path = os.path.join(main_path, filename)
    # template.save(filled_path)
    # print("Done with MIGRATION CERTIFIACTE.docx")
    
    # completed Migration certificate Filling 
    
    # if(certificate == "RESULT CANCELLATION CERTIFICATE"):
    # change because of some reason 7 ki jaga 
    if certificate_name_code == '02000280':
        
        
        # # For Application form for migration request

        # AMR_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Name, [reg_no_])


        # Name_AMR = cursor.fetchone()[0]

        # AMR_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_fname, [reg_no_])

        # fname_AMR = cursor.fetchone()[0]

        # AMR_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        # cursor_2.execute(AMR_roll_no, [reg_no_])

        # roll_no_AMR = cursor_2.fetchone()[0]

        # AMR_reg_no = '''Select reg_no from ZReg where fname = ?'''
        # cursor.execute(AMR_reg_no, [fname_AMR])

        # reg_no_AMR = cursor.fetchone()[0]

        # AMR_Year = '''Select year from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Year, [reg_no_])

        # Year_AMR = cursor.fetchone()[0]

        # # For Migration Certificate    
# .....
        R_C_A_X_roll_no = '''select roll_no from ZLedger where reg_no = ?
'''
        cursor_2.execute(R_C_A_X_roll_no, [reg_no_])

        roll_no_R_C_A_X = cursor_2.fetchone()[0]
        print(roll_no_R_C_A_X)
        
        R_C_A_X_group = '''Select grp_code from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  '''
        cursor_2.execute(R_C_A_X_group, [roll_no_R_C_A_X])

        R_C_A_X_group = cursor_2.fetchone()[0]
        if R_C_A_X_group == '1':
            R_C_A_X_group = "SCIENCE"
        elif R_C_A_X_group == '2':
            R_C_A_X_group = "HUMANITIES"
        elif R_C_A_X_group =='3':
            R_C_A_X_group ="TECHNICAL"
        elif R_C_A_X_group == '4':
            R_C_A_X_group = "MATRIC TECH"
            

        R_C_A_X_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_Name, [reg_no_])

        Name_R_C_A_X = cursor.fetchone()[0]  



        R_C_A_X_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_fname, [reg_no_])

        fname_R_C_A_X = cursor.fetchone()[0]

        R_C_A_X_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor.execute(R_C_A_X_reg_no, [fname_R_C_A_X])

        reg_no_R_C_A_X = cursor.fetchone()[0]



        R_C_A_X_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_Year, [reg_no_])

        Year_R_C_A_X = cursor.fetchone()[0]

        R_C_A_X_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor.execute(R_C_A_X_Institution, [reg_no_])
        Institution_R_C_A_X  = cursor.fetchone()[0]


        # total_marks = '''Select marks_obt from ZLedgerII where reg_no = ?  ''' 
        # cursor_2.execute(total_marks, [reg_no_])
        # Institution_M_C = cursor_2.fetchone()[0]

        # R_C_A_X_sts = '''Select pass_fail_status from ZLedgerII where reg = ?'''
        # cursor_2.execute(R_C_A_X_sts, [reg_no_])
        # R_C_A_X_sts= cursor_2.fetchone()[0]

        R_C_A_X_eng = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?   and sub_code = '001'
'''
        cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X ])
        R_C_A_X_eng= cursor_2.fetchone()[0]
        print(R_C_A_X_eng)


        R_C_A_X_urdu = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '002' '''
        cursor_2.execute(R_C_A_X_urdu, [roll_no_R_C_A_X ])
        R_C_A_X_urdu= cursor_2.fetchone()[0]
        try:
            R_C_A_X_isl = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '008' '''
            cursor_2.execute(R_C_A_X_isl, [roll_no_R_C_A_X ])
            R_C_A_X_isl= cursor_2.fetchone()[0]
        except:
            R_C_A_X_isl = 0
            

        try:
            R_C_A_X_pakS = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '106' '''
            cursor_2.execute(R_C_A_X_pakS, [roll_no_R_C_A_X ])
            R_C_A_X_pakS= cursor_2.fetchone()[0]
        except:
            R_C_A_X_pakS = 0
        try:
            R_C_A_X_math = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '010' '''
            cursor_2.execute(R_C_A_X_math, [roll_no_R_C_A_X ])
            R_C_A_X_math= cursor_2.fetchone()[0]
        except:
            R_C_A_X_math = 0
        print(R_C_A_X_math)

        try:
            R_C_A_X_phy = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy, [roll_no_R_C_A_X ])
            R_C_A_X_phy= cursor_2.fetchone()[0]
        except:
            R_C_A_X_phy = 0
        
        R_C_A_X_phy_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
        cursor_2.execute(R_C_A_X_phy_p, [roll_no_R_C_A_X ])
        R_C_A_X_phy_p= cursor_2.fetchone()[0]

        R_C_A_X_che = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
        cursor_2.execute(R_C_A_X_che, [roll_no_R_C_A_X ])
        R_C_A_X_che= cursor_2.fetchone()[0]
        
        
        R_C_A_X_che_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
        cursor_2.execute(R_C_A_X_che_p, [roll_no_R_C_A_X ])
        R_C_A_X_che_p= cursor_2.fetchone()[0]
        try:
            R_C_A_X_bio = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio, [roll_no_R_C_A_X ])
            R_C_A_X_bio= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio = 0
            
        try:
            R_C_A_X_bio_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio_p= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio_p = 0

        try:
            R_C_A_X_obt = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '013' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_obt, [reg_no_])
            R_C_A_X_obt= cursor_2.fetchone()[0]
        except:
            R_C_A_X_obt = 0
            
        try:
            R_C_A_X_obt_2 = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '013' and app_sts ='2' and  ZLedger.reg_no = ?'''
            cursor_2.execute(R_C_A_X_obt_2, [reg_no_])
            R_C_A_X_obt_2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_obt_2 = 0
            
            
        R_C_A_X_eng2 = ''' Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '001' '''
        cursor_2.execute(R_C_A_X_eng2, [roll_no_R_C_A_X ])
        R_C_A_X_eng2= cursor_2.fetchone()[0]
        print(R_C_A_X_eng2)


        R_C_A_X_urdu2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '002' '''
        cursor_2.execute(R_C_A_X_urdu2, [roll_no_R_C_A_X ])
        R_C_A_X_urdu2= cursor_2.fetchone()[0] 
        try:
            R_C_A_X_isl2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '008' '''
            cursor_2.execute(R_C_A_X_isl2, [roll_no_R_C_A_X ])
            R_C_A_X_isl2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_isl2 = 0
            

        try:
            R_C_A_X_pakS2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '106' '''
            cursor_2.execute(R_C_A_X_pakS2, [roll_no_R_C_A_X ])
            R_C_A_X_pakS2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_pakS2 =0

        try:
            R_C_A_X_math2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '010' '''
            cursor_2.execute(R_C_A_X_math2, [roll_no_R_C_A_X ])
            R_C_A_X_math2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_math2 = 0
        print(R_C_A_X_math2)

        R_C_A_X_phy2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
        cursor_2.execute(R_C_A_X_phy2, [roll_no_R_C_A_X ])
        R_C_A_X_phy2= cursor_2.fetchone()[0]
        
        try:
            R_C_A_X_phy2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
            cursor_2.execute(R_C_A_X_phy2_p, [roll_no_R_C_A_X ])
            R_C_A_X_phy2_p= cursor_2.fetchone()[0]
            
        except:
            R_C_A_X_phy2_p = 0

        try:
            R_C_A_X_che2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_2.execute(R_C_A_X_che2, [roll_no_R_C_A_X ])
            R_C_A_X_che2= cursor_2.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0

        try:
            R_C_A_X_che2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_2.execute(R_C_A_X_che2_p, [roll_no_R_C_A_X ])
            R_C_A_X_che2_p= cursor_2.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0
            
        try:
            R_C_A_X_bio2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio2, [roll_no_R_C_A_X ])
            R_C_A_X_bio2= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
        try:
            R_C_A_X_bio2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_2.execute(R_C_A_X_bio2_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio2_p= cursor_2.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
            
        
            

            
        
        

        # R_C_A_X_total = '''SELECT total FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_total, [reg_no_])
        R_C_A_X_total= 1100
        
        # R_C_A_X_idf = '''SELECT id_mark FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_idf, [reg_no_])
        R_C_A_X_idf="Nothing"
        
        R_C_A_X_sts_e = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_e, [reg_no_])
        R_C_A_X_sts_e= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_e== '1'):
            R_C_A_X_sts_e = "Pass"
        
        else:
            
            R_C_A_X_sts_e = "Fail"
            
        R_C_A_X_sts_u = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '002' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_u, [reg_no_])
        R_C_A_X_sts_u= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_u== '1'):
            R_C_A_X_sts_u = "Pass"
        
        else:
            
            R_C_A_X_sts_u = "Fail"
            
        R_C_A_X_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        cursor.execute(R_C_A_X_DOB, [reg_no_])

        DOB_R_C_A_X = cursor.fetchone()[0]
        
        
        R_C_A_X_sts_i = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_i, [reg_no_])
        R_C_A_X_sts_i= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_i== '1'):
            R_C_A_X_sts_i = "Pass"
        
        else:
            
            R_C_A_X_sts_i = "Fail"
            
        R_C_A_X_sts_pk = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_2.execute(R_C_A_X_sts_pk, [reg_no_])
        R_C_A_X_sts_pk= cursor_2.fetchone()[0]
        if(R_C_A_X_sts_pk== '1'):
            R_C_A_X_sts_pk = "Pass"
        
        else:
            
            R_C_A_X_sts_pk = "Fail"
            
            
        


        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]

        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]











        template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION SSC both_temp.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {
                    'Roll_no' : None,
                    'Group' : None,
                    'Registration_no' : None,
                    'Certificate_no' : None,
                    'Candidate_name':None,
                    'Father_name': None,
                    'Institution': None,
                    'sts_e':None,
                    'sts_u':None,
                    'sts_i':None,
                    'sts_pk':None,
                    'sts_m':None,
                    'sts_ph':None,
                    'sts_ch':None,
                    'sts_bio':None,
                    'sts_ph_p':None,
                    'sts_ch_p':None,
                    'sts_bio_p':None,
                    
                    'R_C_A_X_eng2':None,
                    'R_C_A_X_urdu2':None,
                    'R_C_A_X_isl2':None,
                    'R_C_A_X_paskS2':None,
                    'R_C_A_X_math2':None,
                    'R_C_A_X_phy2':None,
                    'R_C_A_X_che2':None,
                    'R_C_A_X_bio2':None,
                    
                    'R_C_A_X_phy2_p':None,
                    'R_C_A_X_che2_p':None,
                    'R_C_A_X_bio2_p':None,
                    
                    'R_C_A_X_eng':None,
                    'R_C_A_X_urdu':None,
                    'R_C_A_X_isl':None,
                    'R_C_A_X_paskS':None,
                    'R_C_A_X_math':None,
                    'R_C_A_X_phy':None,
                    'R_C_A_X_che':None,
                    'R_C_A_X_bio':None,
                    
                    'R_C_A_X_eng_T':None,
                    'R_C_A_X_urdu_T':None,
                    'R_C_A_X_isl_T':None,
                    'R_C_A_X_paskS_T':None,
                    'R_C_A_X_math_T':None,
                    'R_C_A_X_phy_T':None,
                    'R_C_A_X_che_T':None,
                    'R_C_A_X_bio_T':None,
                    
                    'R_C_A_X_phy_p':None,
                    'R_C_A_X_che_p':None,
                    'R_C_A_X_bio_p':None,
                    
                    'total_marks':None,
                    'idf':None,
                    'year':None,
                    
                    
                    'total_mark_obt':None,
                    'total_mark_obt_2':None,
                    
                    
                    
                    
                    
                    'Dated' : None,
                    'DOB': None,
                    'Examination': None,
                    'Session': None,
                    'year': None,
                    'Status': None,
                    
                    
                    
                    }




        to_fill_in['Roll_no'] = roll_no_R_C_A_X
        to_fill_in['Group'] = R_C_A_X_group
        to_fill_in['Registration_no'] = reg_no_R_C_A_X
        to_fill_in['Certificate_no'] = 1223344
        to_fill_in['Candidate_name'] =Name_R_C_A_X
        to_fill_in['Father_name'] = fname_R_C_A_X
        to_fill_in['Institution'] = Institution_R_C_A_X 
        to_fill_in['total_mark_obt_2'] = R_C_A_X_obt_2
        R_C_A_X_obt = int(R_C_A_X_obt_2)+int(R_C_A_X_obt)

        to_fill_in['total_mark_obt'] = 845
        # to_fill_in['year']= Year
        
        
        to_fill_in['Dated']=  Dated
        
        to_fill_in['sts_e']=R_C_A_X_sts_e
        to_fill_in['sts_u']=R_C_A_X_sts_e
        to_fill_in['sts_i']=R_C_A_X_sts_e
        to_fill_in['sts_pk']=R_C_A_X_sts_e
        to_fill_in['sts_m']=R_C_A_X_sts_e
        to_fill_in['sts_ph']=R_C_A_X_sts_e
        to_fill_in['sts_ch']=R_C_A_X_sts_e
        to_fill_in['sts_bio']=R_C_A_X_sts_e
        
        
        to_fill_in['R_C_A_X_eng']= R_C_A_X_eng 
        to_fill_in['R_C_A_X_urdu']=R_C_A_X_urdu
        to_fill_in['R_C_A_X_isl']= R_C_A_X_isl
        to_fill_in['R_C_A_X_pakS']= R_C_A_X_pakS
        to_fill_in['R_C_A_X_math']=R_C_A_X_math
        to_fill_in['R_C_A_X_phy']= R_C_A_X_phy
        to_fill_in['R_C_A_X_che']=R_C_A_X_che
        to_fill_in['R_C_A_X_bio']= R_C_A_X_bio
        
        to_fill_in['R_C_A_X_phy_p']= R_C_A_X_phy_p
        to_fill_in['R_C_A_X_che_p']=R_C_A_X_che_p
        to_fill_in['R_C_A_X_bio_p']= R_C_A_X_bio_p
        
        to_fill_in['R_C_A_X_eng2']= R_C_A_X_eng2
        to_fill_in['R_C_A_X_urdu2']=R_C_A_X_urdu2
        to_fill_in['R_C_A_X_isl2']= R_C_A_X_isl2
        to_fill_in['R_C_A_X_pakS2']= R_C_A_X_pakS2
        to_fill_in['R_C_A_X_math2']=R_C_A_X_math2
        to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        
        #  addition for the total Values
        
        total_eng = int(R_C_A_X_eng2)+int(R_C_A_X_eng)
        total_urdu =  int(R_C_A_X_urdu2)+int(R_C_A_X_urdu)
        total_isl = int(R_C_A_X_isl2)+int(R_C_A_X_isl)
        total_paks = int(R_C_A_X_pakS2)+int(R_C_A_X_pakS)
        total_math = int(R_C_A_X_math2)+int(R_C_A_X_math)
        total_phy = int(R_C_A_X_phy2)+int(R_C_A_X_phy)
        total_che =int(R_C_A_X_che2)+int(R_C_A_X_che)
        total_bio = int(R_C_A_X_bio2)+int(R_C_A_X_bio)
        
        # end
        to_fill_in['R_C_A_X_eng_T']= total_eng 
        to_fill_in['R_C_A_X_urdu_T']=total_urdu
        to_fill_in['R_C_A_X_isl_T']= total_isl
        to_fill_in['R_C_A_X_pakS_T']= total_paks
        to_fill_in['R_C_A_X_math_T']=total_math
        to_fill_in['R_C_A_X_phy_T']=  total_phy
        to_fill_in['R_C_A_X_che_T']=total_che
        to_fill_in['R_C_A_X_bio_T']= total_bio
        
        to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        to_fill_in['R_C_A_X_phy2_p']= R_C_A_X_phy2_p
        to_fill_in['R_C_A_X_che2_p']=R_C_A_X_che2_p
        to_fill_in['R_C_A_X_bio2_p']= R_C_A_X_bio2_p
        
        
        to_fill_in['total_marks']= R_C_A_X_total


        to_fill_in['Session'] = "Final Session"
        to_fill_in['year'] =  2002
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_R_C_A_X
        to_fill_in['DOB'] = DOB_R_C_A_X 
        to_fill_in['idf']= R_C_A_X_idf
            
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'Result Card ANNUAL EXAMINATION 2020.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Result Card ANNUAL EXAMINATION 2020.docx")
        doc = Document('Result Card ANNUAL EXAMINATION 2020.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_2.execute(data, [reg_no_])
        data = cursor_2.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('Result Card ANNUAL EXAMINATION 2020.docx')

        convert("Result Card ANNUAL EXAMINATION 2020.docx", "Result Card ANNUAL EXAMINATION 2020.pdf")

        images = convert_from_path("Result Card ANNUAL EXAMINATION 2020.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'Result Card ANNUAL EXAMINATION 2020'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        IMAGE_PATH = 'sir-ok.jpg'
        # WIDTH, HEIGTH = 600, 600

        # wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
    

        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        
        # Add image file
        # bg = PhotoImage(file = "sir-ok.jpg")
        
        # # Show image using label
        # label1 = Label( root, image = bg)
        # label1.place(x = 0, y = 0)
        
        
        

        # # frame = Frame(wind, width=600, height=400)
        # frame.pack()
        # frame.place(anchor='center', relx=0.5, rely=0.5)

        # # Create an object of tkinter ImageTk
        # img = ImageTk.PhotoImage(Image.open("sir-ok.jpg"))

        # # Create a Label Widget to display the text or Image
        # label = Label(frame, image = img)
        # label.pack()


        #Load an image in the script
        img= (Image.open("Result Card ANNUAL EXAMINATION 2020.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

        


        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),  OpenScanDocumentFile()  ])
        btn_.place(x = 45,
                y=500,
                width=40,
                height=40
                
                        
                        ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(),OpenScanDocumentFile() ])
        btn2_.place(x = 130,
                y=500,
                width=40,
                height=40
                
                )
        
        
        
        

        wind.mainloop()
        
        # completed Result Card Annual Examination Filling 
        
        
    if certificate_name_code == '03000316':
        
        
        # # For Application form for migration request

        # AMR_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Name, [reg_no_])


        # Name_AMR = cursor.fetchone()[0]

        # AMR_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_fname, [reg_no_])

        # fname_AMR = cursor.fetchone()[0]

        # AMR_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        # cursor_2.execute(AMR_roll_no, [reg_no_])

        # roll_no_AMR = cursor_2.fetchone()[0]

        # AMR_reg_no = '''Select reg_no from ZReg where fname = ?'''
        # cursor.execute(AMR_reg_no, [fname_AMR])

        # reg_no_AMR = cursor.fetchone()[0]

        # AMR_Year = '''Select year from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Year, [reg_no_])

        # Year_AMR = cursor.fetchone()[0]

        # # For Migration Certificate    
# .....
        R_C_A_X_roll_no = '''select roll_no from ZLedger where reg_no = ?
'''
        cursor_4.execute(R_C_A_X_roll_no, [reg_no_])

        roll_no_R_C_A_X = cursor_4.fetchone()[0]
        print(roll_no_R_C_A_X)
        
        R_C_A_X_group = '''Select grp_code from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  '''
        cursor_4.execute(R_C_A_X_group, [roll_no_R_C_A_X])

        R_C_A_X_group = cursor_4.fetchone()[0]
        if R_C_A_X_group == '1':
            R_C_A_X_group = "SCIENCE"
        elif R_C_A_X_group == '2':
            R_C_A_X_group = "PRE-Medical"
        elif R_C_A_X_group =='3':
            R_C_A_X_group ="TECHNICAL"
        elif R_C_A_X_group == '4':
            R_C_A_X_group = "MATRIC TECH"
            

        R_C_A_X_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_Name, [reg_no_])

        Name_R_C_A_X = cursor_3.fetchone()[0]  



        R_C_A_X_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_fname, [reg_no_])

        fname_R_C_A_X = cursor_3.fetchone()[0]

        R_C_A_X_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor_3.execute(R_C_A_X_reg_no, [fname_R_C_A_X])

        reg_no_R_C_A_X = cursor_3.fetchone()[0]



        R_C_A_X_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_Year, [reg_no_])

        Year_R_C_A_X = cursor_3.fetchone()[0]

        R_C_A_X_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor_3.execute(R_C_A_X_Institution, [reg_no_])
        Institution_R_C_A_X  = cursor_3.fetchone()[0]


        # total_marks = '''Select marks_obt from ZLedgerII where reg_no = ?  ''' 
        # cursor_2.execute(total_marks, [reg_no_])
        # Institution_M_C = cursor_2.fetchone()[0]

        # R_C_A_X_sts = '''Select pass_fail_status from ZLedgerII where reg = ?'''
        # cursor_2.execute(R_C_A_X_sts, [reg_no_])
        # R_C_A_X_sts= cursor_2.fetchone()[0]
        
        R_C_A_X_img = 'Untitled.png' 

        R_C_A_X_eng = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '001'
'''
        cursor_4.execute(R_C_A_X_eng, [roll_no_R_C_A_X ])
        R_C_A_X_eng= cursor_4.fetchone()[0]
        print(R_C_A_X_eng)


        R_C_A_X_urdu = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '005' '''
        cursor_4.execute(R_C_A_X_urdu, [roll_no_R_C_A_X ])
        R_C_A_X_urdu= cursor_4.fetchone()[0]
        try:
            R_C_A_X_isl = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
            cursor_4.execute(R_C_A_X_isl, [roll_no_R_C_A_X ])
            R_C_A_X_isl= cursor_4.fetchone()[0]
        except:
            R_C_A_X_isl = 0
            

        try:
            R_C_A_X_pakS = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '106' '''
            cursor_4.execute(R_C_A_X_pakS, [roll_no_R_C_A_X ])
            R_C_A_X_pakS= cursor_4.fetchone()[0]
        
        except:
            R_C_A_X_pakS =0 

        try:
            R_C_A_X_math = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '010' '''
            cursor_4.execute(R_C_A_X_math, [roll_no_R_C_A_X ])
            R_C_A_X_math= cursor_4.fetchone()[0]
        except:
            R_C_A_X_math = 0
        print(R_C_A_X_math)

        try:
            R_C_A_X_phy = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '071' '''
            cursor_4.execute(R_C_A_X_phy, [roll_no_R_C_A_X ])
            R_C_A_X_phy= cursor_4.fetchone()[0]
        except:
            R_C_A_X_phy = 0
        
        try:
            R_C_A_X_phy_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
            cursor_4.execute(R_C_A_X_phy_p, [roll_no_R_C_A_X ])
            R_C_A_X_phy_p= cursor_4.fetchone()[0]
        except:
            R_C_A_X_phy_p = 0
        try:
            R_C_A_X_che = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '073'   '''
            cursor_4.execute(R_C_A_X_che, [roll_no_R_C_A_X ])
            R_C_A_X_che= cursor_4.fetchone()[0]
        except:
            R_C_A_X_che = 0
        
        try:
            R_C_A_X_che_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
            cursor_4.execute(R_C_A_X_che_p, [roll_no_R_C_A_X ])
            R_C_A_X_che_p= cursor_4.fetchone()[0]
        except:
            R_C_A_X_che_p = 0
        try:
            R_C_A_X_bio = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '075' '''
            cursor_4.execute(R_C_A_X_bio, [roll_no_R_C_A_X ])
            R_C_A_X_bio= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio = 0
            
        try:
            R_C_A_X_bio_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_4.execute(R_C_A_X_bio_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio_p= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio_p = 0

        try:
            R_C_A_X_obt = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
            cursor_4.execute(R_C_A_X_obt, [reg_no_])
            R_C_A_X_obt= cursor_4.fetchone()[0]
        except:
            R_C_A_X_obt = 0
            
            
#         R_C_A_X_eng2 = ''' Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '001' '''
#         cursor_4.execute(R_C_A_X_eng2, [roll_no_R_C_A_X ])
#         R_C_A_X_eng2= cursor_4.fetchone()[0]
#         print(R_C_A_X_eng2)


#         R_C_A_X_urdu2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '002' '''
#         cursor_4.execute(R_C_A_X_urdu2, [roll_no_R_C_A_X ])
#         R_C_A_X_urdu2= cursor_4.fetchone()[0] 
#         try:
#             R_C_A_X_isl2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '008' '''
#             cursor_4.execute(R_C_A_X_isl2, [roll_no_R_C_A_X ])
#             R_C_A_X_isl2= cursor_4.fetchone()[0]
#         except:
#             R_C_A_X_isl2 = 0


#         R_C_A_X_pakS2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '106' '''
#         cursor_4.execute(R_C_A_X_pakS2, [roll_no_R_C_A_X ])
#         R_C_A_X_pakS2= cursor_4.fetchone()[0]

#         try:
#             R_C_A_X_math2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '010' '''
#             cursor_4.execute(R_C_A_X_math2, [roll_no_R_C_A_X ])
#             R_C_A_X_math2= cursor_4.fetchone()[0]
#         except:
#             R_C_A_X_math2 = 0
#         print(R_C_A_X_math2)

#         R_C_A_X_phy2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
#         cursor_4.execute(R_C_A_X_phy2, [roll_no_R_C_A_X ])
#         R_C_A_X_phy2= cursor_4.fetchone()[0]
        
#         try:
#             R_C_A_X_phy2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
#             cursor_4.execute(R_C_A_X_phy2_p, [roll_no_R_C_A_X ])
#             R_C_A_X_phy2_p= cursor_4.fetchone()[0]
            
#         except:
#             R_C_A_X_phy2_p = 0

#         try:
#             R_C_A_X_che2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
#             cursor_4.execute(R_C_A_X_che2, [roll_no_R_C_A_X ])
#             R_C_A_X_che2= cursor_4.fetchone()[0]
        
#         except:
#             R_C_A_X_che2_p = 0

#         try:
#             R_C_A_X_che2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
#             cursor_4.execute(R_C_A_X_che2_p, [roll_no_R_C_A_X ])
#             R_C_A_X_che2_p= cursor_4.fetchone()[0]
        
#         except:
#             R_C_A_X_che2_p = 0
            
#         try:
#             R_C_A_X_bio2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
#             cursor_4.execute(R_C_A_X_bio2, [roll_no_R_C_A_X ])
#             R_C_A_X_bio2= cursor_4.fetchone()[0]
#         except:
#             R_C_A_X_bio2_p = 0
#         try:
#             R_C_A_X_bio2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
# ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
# and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
#             cursor_4.execute(R_C_A_X_bio2_p, [roll_no_R_C_A_X ])
#             R_C_A_X_bio2_p= cursor_4.fetchone()[0]
#         except:
#             R_C_A_X_bio2_p = 0
            
        
            

            
        
        

        # R_C_A_X_total = '''SELECT total FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_total, [reg_no_])
        R_C_A_X_total= 505
        
        # R_C_A_X_idf = '''SELECT id_mark FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_idf, [reg_no_])
        R_C_A_X_idf="Nothing"
        
        R_C_A_X_sts_e = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_e, [reg_no_])
        R_C_A_X_sts_e= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_e== '1'):
            R_C_A_X_sts_e = "Pass"
        
        else:
            
            R_C_A_X_sts_e = "Fail"
            
        R_C_A_X_sts_u = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '005' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_u, [reg_no_])
        R_C_A_X_sts_u= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_u== '1'):
            R_C_A_X_sts_u = "Pass"
        
        else:
            
            R_C_A_X_sts_u = "Fail"
            
        # R_C_A_X_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        # cursor_3.execute(R_C_A_X_DOB, [reg_no_])

        # DOB_R_C_A_X = cursor_3.fetchone()[0]
        
        
        try:
            R_C_A_X_sts_i = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '011' and  ZLedger.reg_no = ?'''
            cursor_4.execute(R_C_A_X_sts_i, [reg_no_])
            R_C_A_X_sts_i= cursor_4.fetchone()[0]
        except:
            R_C_A_X_sts_i = 0
        
        if(R_C_A_X_sts_i== '1'):
            R_C_A_X_sts_i = "Pass"
        
        else:
            
            R_C_A_X_sts_i = "Fail"
            
        try:
            R_C_A_X_sts_pk = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
            cursor_4.execute(R_C_A_X_sts_pk, [reg_no_])
            R_C_A_X_sts_pk= cursor_4.fetchone()[0]
        except:
            R_C_A_X_sts_pk = 0
        if(R_C_A_X_sts_pk== '1'):
            R_C_A_X_sts_pk = "Pass"
        
        else:
            
            R_C_A_X_sts_pk = "Fail"
            
            
        


        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]

        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]











        template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION hssc_temp.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')


        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {
                    'Roll_no' : None,
                    'Group' : None,
                    'Registration_no' : None,
                    'Certificate_no' : None,
                    'Candidate_name':None,
                    'Father_name': None,
                    'Institution': None,
                    'sts_e':None,
                    'sts_u':None,
                    'sts_i':None,
                    'sts_pk':None,
                    'sts_m':None,
                    'sts_ph':None,
                    'sts_ch':None,
                    'sts_bio':None,
                    'sts_ph_p':None,
                    'sts_ch_p':None,
                    'sts_bio_p':None,
                    
                    'R_C_A_X_eng2':None,
                    'R_C_A_X_urdu2':None,
                    'R_C_A_X_isl2':None,
                    'R_C_A_X_paskS2':None,
                    'R_C_A_X_math2':None,
                    'R_C_A_X_phy2':None,
                    'R_C_A_X_che2':None,
                    'R_C_A_X_bio2':None,
                    
                    'R_C_A_X_phy2_p':None,
                    'R_C_A_X_che2_p':None,
                    'R_C_A_X_bio2_p':None,
                    'R_C_A_X_img':None,
                    
                    'R_C_A_X_eng':None,
                    'R_C_A_X_urdu':None,
                    'R_C_A_X_isl':None,
                    'R_C_A_X_paskS':None,
                    'R_C_A_X_math':None,
                    'R_C_A_X_phy':None,
                    'R_C_A_X_che':None,
                    'R_C_A_X_bio':None,
                    
                    'R_C_A_X_eng_T':None,
                    'R_C_A_X_urdu_T':None,
                    'R_C_A_X_isl_T':None,
                    'R_C_A_X_paskS_T':None,
                    'R_C_A_X_math_T':None,
                    'R_C_A_X_phy_T':None,
                    'R_C_A_X_che_T':None,
                    'R_C_A_X_bio_T':None,
                    
                    'R_C_A_X_phy_p':None,
                    'R_C_A_X_che_p':None,
                    'R_C_A_X_bio_p':None,
                    
                    'total_marks':None,
                    'idf':None,
                    'year':None,
                    
                    
                    
                    'total_mark_obt':None,
                    
                    
                    
                    'Dated' : None,
                    'DOB': None,
                    'Examination': None,
                    'Session': None,
                    'Year': None,
                    'Status': None,
                    
                    
                    
                    }




        to_fill_in['Roll_no'] = roll_no_R_C_A_X
        to_fill_in['Group'] = R_C_A_X_group
        to_fill_in['Registration_no'] = reg_no_
        to_fill_in['Certificate_no'] = 1223344
        to_fill_in['Candidate_name'] =Name_R_C_A_X
        to_fill_in['Father_name'] = fname_R_C_A_X
        to_fill_in['Institution'] = Institution_R_C_A_X 

        to_fill_in['total_mark_obt'] = R_C_A_X_obt
        to_fill_in['Dated']=  Dated
        to_fill_in['R_C_A_X_img']= R_C_A_X_img
        
        to_fill_in['sts_e']=R_C_A_X_sts_e
        to_fill_in['sts_u']=R_C_A_X_sts_e
        to_fill_in['sts_i']=R_C_A_X_sts_e
        to_fill_in['sts_pk']=R_C_A_X_sts_e
        to_fill_in['sts_m']=R_C_A_X_sts_e
        to_fill_in['sts_ph']=R_C_A_X_sts_e
        to_fill_in['sts_ch']=R_C_A_X_sts_e
        to_fill_in['sts_bio']=R_C_A_X_sts_e
        
        
        to_fill_in['R_C_A_X_eng']= R_C_A_X_eng 
        to_fill_in['R_C_A_X_urdu']=R_C_A_X_urdu
        to_fill_in['R_C_A_X_isl']= R_C_A_X_isl
        to_fill_in['R_C_A_X_pakS']= R_C_A_X_pakS
        to_fill_in['R_C_A_X_math']=R_C_A_X_math
        to_fill_in['R_C_A_X_phy']= R_C_A_X_phy
        to_fill_in['R_C_A_X_che']=R_C_A_X_che
        to_fill_in['R_C_A_X_bio']= R_C_A_X_bio
        
        to_fill_in['R_C_A_X_phy_p']= R_C_A_X_phy_p
        to_fill_in['R_C_A_X_che_p']=R_C_A_X_che_p
        to_fill_in['R_C_A_X_bio_p']= R_C_A_X_bio_p
        
        # to_fill_in['Year']=Year
        
        
        # to_fill_in['R_C_A_X_eng2']= R_C_A_X_eng2
        # to_fill_in['R_C_A_X_urdu2']=R_C_A_X_urdu2
        # to_fill_in['R_C_A_X_isl2']= R_C_A_X_isl2
        # to_fill_in['R_C_A_X_pakS2']= R_C_A_X_pakS2
        # to_fill_in['R_C_A_X_math2']=R_C_A_X_math2
        # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        
        #  addition for the total Values
        
        # total_eng = int(R_C_A_X_eng2)+int(R_C_A_X_eng)
        # total_urdu =  int(R_C_A_X_urdu2)+int(R_C_A_X_urdu)
        # total_isl = int(R_C_A_X_isl2)+int(R_C_A_X_isl)
        # total_paks = int(R_C_A_X_pakS2)+int(R_C_A_X_pakS)
        # total_math = int(R_C_A_X_math2)+int(R_C_A_X_math)
        # total_phy = int(R_C_A_X_phy2)+int(R_C_A_X_phy)
        # total_che =int(R_C_A_X_che2)+int(R_C_A_X_che)
        # total_bio = int(R_C_A_X_bio2)+int(R_C_A_X_bio)
        
        # end
        # to_fill_in['R_C_A_X_eng_T']= total_eng 
        # to_fill_in['R_C_A_X_urdu_T']=total_urdu
        # to_fill_in['R_C_A_X_isl_T']= total_isl
        # to_fill_in['R_C_A_X_pakS_T']= total_paks
        # to_fill_in['R_C_A_X_math_T']=total_math
        # to_fill_in['R_C_A_X_phy_T']=  total_phy
        # to_fill_in['R_C_A_X_che_T']=total_che
        # to_fill_in['R_C_A_X_bio_T']= total_bio
        
        # to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        # to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        # to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        # to_fill_in['R_C_A_X_phy2_p']= R_C_A_X_phy2_p
        # to_fill_in['R_C_A_X_che2_p']=R_C_A_X_che2_p
        # to_fill_in['R_C_A_X_bio2_p']= R_C_A_X_bio2_p
        
        
        to_fill_in['total_marks']= R_C_A_X_total


        to_fill_in['Session'] = "Final Session"
        # to_fill_in['year'] =  2002
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_R_C_A_X
        # to_fill_in['DOB'] = DOB_R_C_A_X 
        to_fill_in['idf']= R_C_A_X_idf
        # to_fill_in['year']= year    
        
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'Result Card ANNUAL EXAMINATION HSSC.docx'
    
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Result Card ANNUAL EXAMINATION HSSC.docx")
        doc = Document('Result Card ANNUAL EXAMINATION HSSC.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_4.execute(data, [reg_no_])
        data = cursor_4.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('Result Card ANNUAL EXAMINATION HSSC.docx')

        convert("Result Card ANNUAL EXAMINATION HSSC.docx", "Result Card ANNUAL EXAMINATION HSSC.pdf")
        

        images = convert_from_path("Result Card ANNUAL EXAMINATION HSSC.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'Result Card ANNUAL EXAMINATION HSSC'+'.png'
            image.save(fname, "PNG")
            

            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()
        
        IMAGE_PATH = 'sir-ok.jpg'
# WIDTH, HEIGTH = 600, 600

# wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        # img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        # lbl = tk.Label(wind, image=img)
        # lbl.img = img  # Keep a reference in case this code put is in a function.
        # lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        # #Create a canvas
        # canvas= Canvas(wind, width= 900, height= 900)
        # canvas.pack()
        
        
        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        # canvas= Canvas(wind, width= 900, height= 900)
        # canvas.pack()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        

        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        
        # Add image file
        # bg = PhotoImage(file = "sir-ok.jpg")
        
        # # Show image using label
        # label1 = Label( root, image = bg)
        # label1.place(x = 0, y = 0)
        
        
        

        # # frame = Frame(wind, width=600, height=400)
        # frame.pack()
        # frame.place(anchor='center', relx=0.5, rely=0.5)

        # # Create an object of tkinter ImageTk
        # img = ImageTk.PhotoImage(Image.open("sir-ok.jpg"))

        # # Create a Label Widget to display the text or Image
        # label = Label(frame, image = img)
        # label.pack()


        #Load an image in the script
        img= (Image.open("Result Card ANNUAL EXAMINATION HSSC.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

    


        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),OpenScanDocumentFile() ])
        btn_.place(x = 45,
                    y=500,
                    width=40,
                    height=40
                    
                            
                            ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(), OpenScanDocumentFile()])
        btn2_.place(x = 130,
                    y=500,
                    width=40,
                    height=40
                    
                    ) 
        
        
        
        

        wind.mainloop()
    
    # change for some reason 7 ki jaga 6 karna 
    if certificate_name_code == '03000380':
        
        
        # # For Application form for migration request

        # AMR_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Name, [reg_no_])


        # Name_AMR = cursor.fetchone()[0]

        # AMR_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_fname, [reg_no_])

        # fname_AMR = cursor.fetchone()[0]

        # AMR_roll_no = '''select roll_no from ZLedgerII where reg_no = ?'''
        # cursor_2.execute(AMR_roll_no, [reg_no_])

        # roll_no_AMR = cursor_2.fetchone()[0]

        # AMR_reg_no = '''Select reg_no from ZReg where fname = ?'''
        # cursor.execute(AMR_reg_no, [fname_AMR])

        # reg_no_AMR = cursor.fetchone()[0]

        # AMR_Year = '''Select year from ZReg WHERE reg_no = ?'''
        # cursor.execute(AMR_Year, [reg_no_])

        # Year_AMR = cursor.fetchone()[0]

        # # For Migration Certificate    
# .....
        R_C_A_X_roll_no = '''select roll_no from ZLedger where reg_no = ?
'''
        cursor_4.execute(R_C_A_X_roll_no, [reg_no_])

        roll_no_R_C_A_X = cursor_4.fetchone()[0]
        print(roll_no_R_C_A_X)
        
        R_C_A_X_group = '''Select grp_code from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  '''
        cursor_4.execute(R_C_A_X_group, [roll_no_R_C_A_X])

        R_C_A_X_group = cursor_4.fetchone()[0]
        if R_C_A_X_group == '1':
            R_C_A_X_group = "SCIENCE"
        elif R_C_A_X_group == '2':
            R_C_A_X_group = "HUMANITIES"
        elif R_C_A_X_group =='3':
            R_C_A_X_group ="TECHNICAL"
        elif R_C_A_X_group == '4':
            R_C_A_X_group = "MATRIC TECH"
            

        R_C_A_X_Name = '''SELECT name from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_Name, [reg_no_])

        Name_R_C_A_X = cursor_3.fetchone()[0]  



        R_C_A_X_fname = '''Select fname from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_fname, [reg_no_])

        fname_R_C_A_X = cursor_3.fetchone()[0]

        R_C_A_X_reg_no = '''Select reg_no from ZReg where fname = ?'''
        cursor_3.execute(R_C_A_X_reg_no, [fname_R_C_A_X])

        reg_no_R_C_A_X = cursor_3.fetchone()[0]



        R_C_A_X_Year = '''Select year from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_Year, [reg_no_])

        Year_R_C_A_X = cursor_3.fetchone()[0]

        R_C_A_X_Institution = '''select inst_desc from ZReg where reg_no = ?'''
        cursor_3.execute(R_C_A_X_Institution, [reg_no_])
        Institution_R_C_A_X  = cursor_3.fetchone()[0]


        # total_marks = '''Select marks_obt from ZLedgerII where reg_no = ?  ''' 
        # cursor_2.execute(total_marks, [reg_no_])
        # Institution_M_C = cursor_2.fetchone()[0]

        # R_C_A_X_sts = '''Select pass_fail_status from ZLedgerII where reg = ?'''
        # cursor_2.execute(R_C_A_X_sts, [reg_no_])
        # R_C_A_X_sts= cursor_2.fetchone()[0]

        R_C_A_X_eng = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '001'
'''
        cursor_4.execute(R_C_A_X_eng, [roll_no_R_C_A_X ])
        R_C_A_X_eng= cursor_4.fetchone()[0]
        print(R_C_A_X_eng)


        R_C_A_X_urdu = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '002' '''
        cursor_4.execute(R_C_A_X_urdu, [roll_no_R_C_A_X ])
        R_C_A_X_urdu= cursor_4.fetchone()[0]
        try:
            R_C_A_X_isl = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '008' '''
            cursor_4.execute(R_C_A_X_isl, [roll_no_R_C_A_X ])
            R_C_A_X_isl= cursor_4.fetchone()[0]
        except:
            R_C_A_X_isl = 0
            

        R_C_A_X_pakS = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '106' '''
        cursor_4.execute(R_C_A_X_pakS, [roll_no_R_C_A_X ])
        R_C_A_X_pakS= cursor_4.fetchone()[0]

        try:
            R_C_A_X_math = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '010' '''
            cursor_4.execute(R_C_A_X_math, [roll_no_R_C_A_X ])
            R_C_A_X_math= cursor_4.fetchone()[0]
        except:
            R_C_A_X_math = 0
        print(R_C_A_X_math)

        R_C_A_X_phy = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
        cursor_4.execute(R_C_A_X_phy, [roll_no_R_C_A_X ])
        R_C_A_X_phy= cursor_4.fetchone()[0]
        
        R_C_A_X_phy_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '011' '''
        cursor_4.execute(R_C_A_X_phy_p, [roll_no_R_C_A_X ])
        R_C_A_X_phy_p= cursor_4.fetchone()[0]

        R_C_A_X_che = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
        cursor_4.execute(R_C_A_X_che, [roll_no_R_C_A_X ])
        R_C_A_X_che= cursor_4.fetchone()[0]
        
        
        R_C_A_X_che_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '012'   '''
        cursor_4.execute(R_C_A_X_che_p, [roll_no_R_C_A_X ])
        R_C_A_X_che_p= cursor_4.fetchone()[0]
        try:
            R_C_A_X_bio = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_4.execute(R_C_A_X_bio, [roll_no_R_C_A_X ])
            R_C_A_X_bio= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio = 0
            
        try:
            R_C_A_X_bio_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ?  and sub_code = '013' '''
            cursor_4.execute(R_C_A_X_bio_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio_p= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio_p = 0

        try:
            R_C_A_X_obt = '''Select marks_obt from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '013' and  ZLedger.reg_no = ?'''
            cursor_4.execute(R_C_A_X_obt, [reg_no_])
            R_C_A_X_obt= cursor_4.fetchone()[0]
        except:
            R_C_A_X_obt = 0
            
            
        R_C_A_X_eng2 = ''' Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '001' '''
        cursor_4.execute(R_C_A_X_eng2, [roll_no_R_C_A_X ])
        R_C_A_X_eng2= cursor_4.fetchone()[0]
        print(R_C_A_X_eng2)


        R_C_A_X_urdu2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '002' '''
        cursor_4.execute(R_C_A_X_urdu2, [roll_no_R_C_A_X ])
        R_C_A_X_urdu2= cursor_4.fetchone()[0] 
        try:
            R_C_A_X_isl2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '008' '''
            cursor_4.execute(R_C_A_X_isl2, [roll_no_R_C_A_X ])
            R_C_A_X_isl2= cursor_4.fetchone()[0]
        except:
            R_C_A_X_isl2 = 0
            

        R_C_A_X_pakS2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '106' '''
        cursor_4.execute(R_C_A_X_pakS2, [roll_no_R_C_A_X ])
        R_C_A_X_pakS2= cursor_4.fetchone()[0]

        try:
            R_C_A_X_math2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '010' '''
            cursor_4.execute(R_C_A_X_math2, [roll_no_R_C_A_X ])
            R_C_A_X_math2= cursor_4.fetchone()[0]
        except:
            R_C_A_X_math2 = 0
        print(R_C_A_X_math2)

        R_C_A_X_phy2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
        cursor_4.execute(R_C_A_X_phy2, [roll_no_R_C_A_X ])
        R_C_A_X_phy2= cursor_4.fetchone()[0]
        
        try:
            R_C_A_X_phy2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '011' '''
            cursor_4.execute(R_C_A_X_phy2_p, [roll_no_R_C_A_X ])
            R_C_A_X_phy2_p= cursor_4.fetchone()[0]
            
        except:
            R_C_A_X_phy2_p = 0

        try:
            R_C_A_X_che2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_4.execute(R_C_A_X_che2, [roll_no_R_C_A_X ])
            R_C_A_X_che2= cursor_4.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0

        try:
            R_C_A_X_che2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '012' '''
            cursor_4.execute(R_C_A_X_che2_p, [roll_no_R_C_A_X ])
            R_C_A_X_che2_p= cursor_4.fetchone()[0]
        
        except:
            R_C_A_X_che2_p = 0
            
        try:
            R_C_A_X_bio2 = '''Select marks_t from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_4.execute(R_C_A_X_bio2, [roll_no_R_C_A_X ])
            R_C_A_X_bio2= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
        try:
            R_C_A_X_bio2_p = '''Select marks_p from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no  where ZLedger.roll_no = ? and app_sts = '1' and sub_code = '013' '''
            cursor_4.execute(R_C_A_X_bio2_p, [roll_no_R_C_A_X ])
            R_C_A_X_bio2_p= cursor_4.fetchone()[0]
        except:
            R_C_A_X_bio2_p = 0
            
        
            

            
        
        

        # R_C_A_X_total = '''SELECT total FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_total, [reg_no_])
        R_C_A_X_total= 1100
        
        # R_C_A_X_idf = '''SELECT id_mark FROM [ssc_LEDGER].[dbo].[ZLedgerII] where reg_no  = ?'''
        # cursor_2.execute(R_C_A_X_idf, [reg_no_])
        R_C_A_X_idf="Nothing"
        
        R_C_A_X_sts_e = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_e, [reg_no_])
        R_C_A_X_sts_e= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_e== '1'):
            R_C_A_X_sts_e = "Pass"
        
        else:
            
            R_C_A_X_sts_e = "Fail"
            
        R_C_A_X_sts_u = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '002' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_u, [reg_no_])
        R_C_A_X_sts_u= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_u== '1'):
            R_C_A_X_sts_u = "Pass"
        
        else:
            
            R_C_A_X_sts_u = "Fail"
            
        R_C_A_X_DOB = '''SELECT dob from ZReg WHERE reg_no = ?'''
        cursor_3.execute(R_C_A_X_DOB, [reg_no_])

        DOB_R_C_A_X = cursor_3.fetchone()[0]
        
        
        R_C_A_X_sts_i = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_i, [reg_no_])
        R_C_A_X_sts_i= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_i== '1'):
            R_C_A_X_sts_i = "Pass"
        
        else:
            
            R_C_A_X_sts_i = "Fail"
            
        R_C_A_X_sts_pk = '''Select pass_fail_sts from ZLedger INNER JOIN ZLedgerSub  on
ZLedger.exam_code  = ZLedgerSub.exam_code and ZLedger.year = ZLedgerSub.year and ZLedger.sess_code = ZLedgerSub.sess_code
and ZLedger.roll_no = ZLedgerSub.roll_no   where ZLedgerSub.sub_code = '001' and  ZLedger.reg_no = ?'''
        cursor_4.execute(R_C_A_X_sts_pk, [reg_no_])
        R_C_A_X_sts_pk= cursor_4.fetchone()[0]
        if(R_C_A_X_sts_pk== '1'):
            R_C_A_X_sts_pk = "Pass"
        
        else:
            
            R_C_A_X_sts_pk = "Fail"
            
            
        


        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]

        # R_C_A_X_eng = '''Select mark_t from ZledgerII where sub_code = 001 and roll no = ?'''
        # cursor_2.execute(R_C_A_X_eng, [roll_no_R_C_A_X])
        # R_C_A_X_eng= cursor_2.fetchone()[0]











        template_path = os.path.join(main_path, 'Result Card ANNUAL EXAMINATION HSSC both_temp.docx')
        workbook_path = os.path.join(main_path, 'Template_data.xlsx')

        workbook = load_workbook(workbook_path)
        template = DocxTemplate(template_path)
        worksheet = workbook["Input"]

        to_fill_in = {
                    'Roll_no' : None,
                    'Group' : None,
                    'Registration_no' : None,
                    'Certificate_no' : None,
                    'Candidate_name':None,
                    'Father_name': None,
                    'Institution': None,
                    'sts_e':None,
                    'sts_u':None,
                    'sts_i':None,
                    'sts_pk':None,
                    'sts_m':None,
                    'sts_ph':None,
                    'sts_ch':None,
                    'sts_bio':None,
                    'sts_ph_p':None,
                    'sts_ch_p':None,
                    'sts_bio_p':None,
                    
                    'R_C_A_X_eng2':None,
                    'R_C_A_X_urdu2':None,
                    'R_C_A_X_isl2':None,
                    'R_C_A_X_paskS2':None,
                    'R_C_A_X_math2':None,
                    'R_C_A_X_phy2':None,
                    'R_C_A_X_che2':None,
                    'R_C_A_X_bio2':None,
                    
                    'R_C_A_X_phy2_p':None,
                    'R_C_A_X_che2_p':None,
                    'R_C_A_X_bio2_p':None,
                    
                    'R_C_A_X_eng':None,
                    'R_C_A_X_urdu':None,
                    'R_C_A_X_isl':None,
                    'R_C_A_X_paskS':None,
                    'R_C_A_X_math':None,
                    'R_C_A_X_phy':None,
                    'R_C_A_X_che':None,
                    'R_C_A_X_bio':None,
                    
                    'R_C_A_X_eng_T':None,
                    'R_C_A_X_urdu_T':None,
                    'R_C_A_X_isl_T':None,
                    'R_C_A_X_paskS_T':None,
                    'R_C_A_X_math_T':None,
                    'R_C_A_X_phy_T':None,
                    'R_C_A_X_che_T':None,
                    'R_C_A_X_bio_T':None,
                    
                    'R_C_A_X_phy_p':None,
                    'R_C_A_X_che_p':None,
                    'R_C_A_X_bio_p':None,
                    
                    'total_marks':None,
                    'idf':None,
                    'year':None,
                    
                    
                    'total_mark_obt':None,
                    
                    
                    
                    'Dated' : None,
                    'DOB': None,
                    'Examination': None,
                    'Session': None,
                    'year': None,
                    'Status': None,
                    
                    
                    
                    }




        to_fill_in['Roll_no'] = roll_no_R_C_A_X
        to_fill_in['Group'] = R_C_A_X_group
        to_fill_in['Registration_no'] = reg_no_R_C_A_X
        to_fill_in['Certificate_no'] = 1223344
        to_fill_in['Candidate_name'] =Name_R_C_A_X
        to_fill_in['Father_name'] = fname_R_C_A_X
        to_fill_in['Institution'] = Institution_R_C_A_X 

        to_fill_in['total_mark_obt'] = 845
        to_fill_in['Dated']=  Dated
        
        to_fill_in['sts_e']=R_C_A_X_sts_e
        to_fill_in['sts_u']=R_C_A_X_sts_e
        to_fill_in['sts_i']=R_C_A_X_sts_e
        to_fill_in['sts_pk']=R_C_A_X_sts_e
        to_fill_in['sts_m']=R_C_A_X_sts_e
        to_fill_in['sts_ph']=R_C_A_X_sts_e
        to_fill_in['sts_ch']=R_C_A_X_sts_e
        to_fill_in['sts_bio']=R_C_A_X_sts_e
        
        
        to_fill_in['R_C_A_X_eng']= R_C_A_X_eng 
        to_fill_in['R_C_A_X_urdu']=R_C_A_X_urdu
        to_fill_in['R_C_A_X_isl']= R_C_A_X_isl
        to_fill_in['R_C_A_X_pakS']= R_C_A_X_pakS
        to_fill_in['R_C_A_X_math']=R_C_A_X_math
        to_fill_in['R_C_A_X_phy']= R_C_A_X_phy
        to_fill_in['R_C_A_X_che']=R_C_A_X_che
        to_fill_in['R_C_A_X_bio']= R_C_A_X_bio
        
        to_fill_in['R_C_A_X_phy_p']= R_C_A_X_phy_p
        to_fill_in['R_C_A_X_che_p']=R_C_A_X_che_p
        to_fill_in['R_C_A_X_bio_p']= R_C_A_X_bio_p
        
        to_fill_in['R_C_A_X_eng2']= R_C_A_X_eng2
        to_fill_in['R_C_A_X_urdu2']=R_C_A_X_urdu2
        to_fill_in['R_C_A_X_isl2']= R_C_A_X_isl2
        to_fill_in['R_C_A_X_pakS2']= R_C_A_X_pakS2
        to_fill_in['R_C_A_X_math2']=R_C_A_X_math2
        to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        
        
        #  addition for the total Values
        
        total_eng = int(R_C_A_X_eng2)+int(R_C_A_X_eng)
        total_urdu =  int(R_C_A_X_urdu2)+int(R_C_A_X_urdu)
        total_isl = int(R_C_A_X_isl2)+int(R_C_A_X_isl)
        total_paks = int(R_C_A_X_pakS2)+int(R_C_A_X_pakS)
        total_math = int(R_C_A_X_math2)+int(R_C_A_X_math)
        total_phy = int(R_C_A_X_phy2)+int(R_C_A_X_phy)
        total_che =int(R_C_A_X_che2)+int(R_C_A_X_che)
        total_bio = int(R_C_A_X_bio2)+int(R_C_A_X_bio)
        
        # end
        to_fill_in['R_C_A_X_eng_T']= total_eng 
        to_fill_in['R_C_A_X_urdu_T']=total_urdu
        to_fill_in['R_C_A_X_isl_T']= total_isl
        to_fill_in['R_C_A_X_pakS_T']= total_paks
        to_fill_in['R_C_A_X_math_T']=total_math
        to_fill_in['R_C_A_X_phy_T']=  total_phy
        to_fill_in['R_C_A_X_che_T']=total_che
        to_fill_in['R_C_A_X_bio_T']= total_bio
        
        to_fill_in['R_C_A_X_phy2']= R_C_A_X_phy2
        to_fill_in['R_C_A_X_che2']=R_C_A_X_che2
        to_fill_in['R_C_A_X_bio2']= R_C_A_X_bio2
        reg_no_
        to_fill_in['R_C_A_X_phy2_p']= R_C_A_X_phy2_p
        to_fill_in['R_C_A_X_che2_p']=R_C_A_X_che2_p
        to_fill_in['R_C_A_X_bio2_p']= R_C_A_X_bio2_p
        
        
        to_fill_in['total_marks']= R_C_A_X_total


        to_fill_in['Session'] = "Final Session"
        to_fill_in['year'] =  2002
        to_fill_in['Status'] = "pass"
        to_fill_in['Institution'] = Institution_R_C_A_X
        to_fill_in['DOB'] = DOB_R_C_A_X 
        to_fill_in['idf']= R_C_A_X_idf
        to_fill_in['year']= year
            
        # Fill in all the keys defined in the word document using the dictionary.
        # The keys in de word document are identified by the {{}}symbols.
        template.render(to_fill_in)
        # Output the file to a docx document.
        filename = 'Result Card ANNUAL EXAMINATION 2020.docx'
        filled_path = os.path.join(main_path, filename)
        template.save(filled_path)
        print("Result Card ANNUAL EXAMINATION 2020.docx")
        doc = Document('Result Card ANNUAL EXAMINATION 2020.docx')
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        # r = p.add_run()
        # r.add_picture('sir.png',width=Inches(4.0), height=Inches(.7))
        # p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        data = ("select pic from [ssc_LEDGER].[dbo].[ZLedger] tab1 left join [ssc_LEDGER_PIC].[dbo].[ZledgerPic] tab2 ON tab1.year  = tab2.year and tab1.sess_code = tab2.sess_code and tab1.roll_no = tab2.roll_no where tab1.reg_no = ? ")
        cursor_2.execute(data, [reg_no_])
        data = cursor_2.fetchone()[0]
        r.add_picture(io.BytesIO(data),width=Inches(1.0), height=Inches(1.0))
        # r.add_picture('sir.png',width=Inches(1.0), height=Inches(1.0))
        doc.save('Result Card ANNUAL EXAMINATION 2020.docx')

        convert("Result Card ANNUAL EXAMINATION 2020.docx", "Result Card ANNUAL EXAMINATION 2020.pdf")

        images = convert_from_path("Result Card ANNUAL EXAMINATION 2020.pdf", 500,poppler_path=r'C:\Program Files\poppler-0.68.0\bin')
        for i, image in enumerate(images):
            fname = 'Result Card ANNUAL EXAMINATION 2020'+'.png'
            image.save(fname, "PNG")
            
        # Python3 program to convert docx to 
            

        #Import the required Libraries

        #Create an instance of tkinter frame
        wind = tk.Toplevel()
        
        IMAGE_PATH = 'sir-ok.jpg'
# WIDTH, HEIGTH = 600, 600

# wind.geometry('{}x{}'.format(WIDTH, HEIGHT))

        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        
        
        img = ImageTk.PhotoImage(Image.open(IMAGE_PATH).resize((1600, 1600), Image.ANTIALIAS))
        lbl = tk.Label(wind, image=img)
        lbl.img = img  # Keep a reference in case this code put is in a function.
        lbl.place(relx=0.5, rely=0.5, anchor='center')  # Place label in center of parent.
        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()

        #Set the geometry of tkinter frame
        wind.geometry("1600x1600")
        wind.attributes('-fullscreen',True)
        

        #Create a canvas
        canvas= Canvas(wind, width= 900, height= 900)
        canvas.pack()
        
        # Add image file
        # bg = PhotoImage(file = "sir-ok.jpg")
        
        # # Show image using label
        # label1 = Label( root, image = bg)
        # label1.place(x = 0, y = 0)
        
        
        

        # # frame = Frame(wind, width=600, height=400)
        # frame.pack()
        # frame.place(anchor='center', relx=0.5, rely=0.5)

        # # Create an object of tkinter ImageTk
        # img = ImageTk.PhotoImage(Image.open("sir-ok.jpg"))

        # # Create a Label Widget to display the text or Image
        # label = Label(frame, image = img)
        # label.pack()


        #Load an image in the script
        img= (Image.open("Result Card ANNUAL EXAMINATION 2020.png"))

        #Resize the Image using resize method
        resized_image= img.resize((900,705), Image.Resampling.LANCZOS)
        new_image= ImageTk.PhotoImage(resized_image)

        #Add image to the Canvas Items
        canvas.create_image(10,10, anchor=NW, image=new_image)

    


        btn_ = Button(wind, text="print", command =  lambda : [Print(),wind.destroy(),OpenScanDocumentFile() ])
        btn_.place(x = 45,
                    y=500,
                    width=40,
                    height=40
                    
                            
                            ) 

        btn2_ = Button(wind, text="No", command = lambda : [wind.destroy(),NotPrint(), OpenScanDocumentFile()])
        btn2_.place(x = 130,
                    y=500,
                    width=40,
                    height=40
                    
                    ) 
    
    
    
    

    wind.mainloop()


# def sleep():
#     time.sleep(8)

def Print():
    
    
    
    
    
    if os.path.exists("MIGRATION CERTIFIACTE.docx"):
        
        
        
        # Ask for file (Which you want to print)
        file_to_print = "MIGRATION CERTIFIACTE.docx"
        
        
        
        if file_to_print:
            
            
            # Print Hard Copy of File
            win32api.ShellExecute(0,              # NULL since it's not associated with a window
                "print",        # execute the "print" verb defined for the file type
                file_to_print,  # path to the document file to print
                None,           #no parameters, since the target is a document file
                ".",            #current directory, same as NULL here
                0)              # SW_HIDE passed to app associated with the file type 
        os.sleep(2)
        if os.path.exists("MIGRATION CERTIFIACTE.pdf"):
            os.remove("MIGRATION CERTIFIACTE.pdf") # one file at a time   
        if os.path.exists("MIGRATION CERTIFIACTE.png"):
            os.remove("MIGRATION CERTIFIACTE.png")
        if os.path.exists("MIGRATION CERTIFIACTE.docx"):
            os.remove("MIGRATION CERTIFIACTE.docx")
            
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.docx"):
        file_to_print = "RESULT CANCELLATION CERTIFICATE.docx"
        
        
        
        if file_to_print:
            
            # Print Hard Copy of File
            win32api.ShellExecute(0,              # NULL since it's not associated with a window
                "print",        # execute the "print" verb defined for the file type
                file_to_print,  # path to the document file to print
                None,           #no parameters, since the target is a document file
                ".",            #current directory, same as NULL here
                0)
        time.slepp(2)
        if os.path.exists("RESULT CANCELLATION CERTIFICATE.pdf"):
            os.remove("RESULT CANCELLATION CERTIFICATE.pdf") # one file at a time   
        if os.path.exists("RESULT CANCELLATION CERTIFICATE.png"):
            os.remove("RESULT CANCELLATION CERTIFICATE.png")
        if os.path.exists("RESULT CANCELLATION CERTIFICATE.docx"):
            os.remove("RESULT CANCELLATION CERTIFICATE.docx")
            
            
            
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.docx"):
        file_to_print = "Result Card ANNUAL EXAMINATION 2020.docx"
        
        
        
        if file_to_print:
            
            # Print Hard Copy of File
            win32api.ShellExecute(0,              # NULL since it's not associated with a window
                "print",        # execute the "print" verb defined for the file type
                file_to_print,  # path to the document file to print
                None,           #no parameters, since the target is a document file
                ".",            #current directory, same as NULL here
                0)
        time.slepp(2)
        if os.path.exists("Result Card ANNUAL EXAMINATION 2020.pdf"):
            os.remove("Result Card ANNUAL EXAMINATION 2020.pdf") # one file at a time   
        if os.path.exists("Result Card ANNUAL EXAMINATION 2020.png"):
            os.remove("Result Card ANNUAL EXAMINATION 2020.png")
        if os.path.exists("Result Card ANNUAL EXAMINATION 2020.docx"):
            os.remove("Result Card ANNUAL EXAMINATION 2020.docx")
            
    # if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.docx"):
    #     file_to_print = "Result Card ANNUAL EXAMINATION HSSC.docx"
        
        
        
    #     if file_to_print:
            
    #         # Print Hard Copy of File
    #         win32api.ShellExecute(0,              # NULL since it's not associated with a window
    #             "print",        # execute the "print" verb defined for the file type
    #             file_to_print,  # path to the document file to print
    #             None,           #no parameters, since the target is a document file
    #             ".",            #current directory, same as NULL here
    #             0)
            
    if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.docx"):
        file_to_print = "Result Card ANNUAL EXAMINATION HSSC.docx"
        
        
        
        if file_to_print:
            
            # Print Hard Copy of File
            win32api.ShellExecute(0,              # NULL since it's not associated with a window
                "print",        # execute the "print" verb defined for the file type
                file_to_print,  # path to the document file to print
                None,           #no parameters, since the target is a document file
                ".",            #current directory, same as NULL here
                0)
        time.slepp(2)
        if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.pdf"):
            os.remove("Result Card ANNUAL EXAMINATION HSSC.pdf") # one file at a time   
        if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.png"):
            os.remove("Result Card ANNUAL EXAMINATION HSSC.png")
        if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.docx"):
            os.remove("Result Card ANNUAL EXAMINATION HSSC.docx")
            
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.docx"):
        file_to_print = "Result Card ANNUAL EXAMINATION SSC.docx"
        
        
        
        if file_to_print:
            
            # Print Hard Copy of File
            win32api.ShellExecute(0,              # NULL since it's not associated with a window
                "print",        # execute the "print" verb defined for the file type
                file_to_print,  # path to the document file to print
                None,           #no parameters, since the target is a document file
                ".",            #current directory, same as NULL here
                0)
            time.sleep(2)
            if os.path.exists("Result Card ANNUAL EXAMINATION SSC.pdf"):
                os.remove("Result Card ANNUAL EXAMINATION SSC.pdf") # one file at a time   
            if os.path.exists("Result Card ANNUAL EXAMINATION SSC.png"):
                os.remove("Result Card ANNUAL EXAMINATION SSC.png")
            if os.path.exists("Result Card ANNUAL EXAMINATION SSC.docx"):
                os.remove("Result Card ANNUAL EXAMINATION SSC.docx")
            
            
            
        
    
# this function is used when we print the document 

    time.sleep(10)
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.docx"):
        os.remove("RESULT CANCELLATION CERTIFICATE.docx") # one file at a time
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.pdf"):
        os.remove("RESULT CANCELLATION CERTIFICATE.pdf") # one file at a time
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.png"):
        os.remove("RESULT CANCELLATION CERTIFICATE.png") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.docx"):
        os.remove("MIGRATION CERTIFIACTE.docx") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.png"):
        os.remove("MIGRATION CERTIFIACTE.png") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.pdf"):
        os.remove("MIGRATION CERTIFIACTE.pdf") # one file at a time 
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.docx"):
            os.remove("Result Card ANNUAL EXAMINATION 2020.docx") # one file at a time
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.png"):
        os.remove("Result Card ANNUAL EXAMINATION 2020.png") # one file at a time
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.pdf"):
        os.remove("Result Card ANNUAL EXAMINATION 2020.pdf") # one file at a time
        
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.pdf"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.pdf") # one file at a time   
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.png"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.png")
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.docx"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.docx")
        
    
    
   
    
                        
    
def thanks():
    thanks = Tk()
    thanks.geometry("1600x1600")



    thanks.after(10000, lambda: thanks.destroy()) # Destroy the widget after 30 seconds





    thanks.mainloop()

    
    
def NotPrint():
    print("not print")
   
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.docx"):
        os.remove("RESULT CANCELLATION CERTIFICATE.docx") # one file at a time
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.pdf"):
        os.remove("RESULT CANCELLATION CERTIFICATE.pdf") # one file at a time
    if os.path.exists("RESULT CANCELLATION CERTIFICATE.png"):
        os.remove("RESULT CANCELLATION CERTIFICATE.png") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.docx"):
        os.remove("MIGRATION CERTIFIACTE.docx") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.png"):
        os.remove("MIGRATION CERTIFIACTE.png") # one file at a time
    if os.path.exists("MIGRATION CERTIFIACTE.pdf"):
        os.remove("MIGRATION CERTIFIACTE.pdf") # one file at a time 
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.docx"):
            os.remove("Result Card ANNUAL EXAMINATION 2020.docx") # one file at a time
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.png"):
        os.remove("Result Card ANNUAL EXAMINATION 2020.png") # one file at a time
    if os.path.exists("Result Card ANNUAL EXAMINATION 2020.pdf"):
        os.remove("Result Card ANNUAL EXAMINATION 2020.pdf") # one file at a time
    if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.pdf"):
        os.remove("Result Card ANNUAL EXAMINATION HSSC.pdf") # one file at a time   
    if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.png"):
        os.remove("Result Card ANNUAL EXAMINATION HSSC.png")
    if os.path.exists("Result Card ANNUAL EXAMINATION HSSC.docx"):
        os.remove("Result Card ANNUAL EXAMINATION HSSC.docx")
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.pdf"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.pdf") # one file at a time   
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.png"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.png")
    if os.path.exists("Result Card ANNUAL EXAMINATION SSC.docx"):
        os.remove("Result Card ANNUAL EXAMINATION SSC.docx")
    
    # Result Card ANNUAL EXAMINATION SSC both_temp.docx
    
# def OpenScanDocumentFile():
#     os.system("python ScanDocument.py")

def OpenScanDocumentFile():
    os.system("python ScanDocument.py")
    

    
    
    

def disable_event():
    pass


root.mainloop()  
    

# def change(root):
#     root.destroy()
#     fun()
    
# def call():
#     root = Tk();
#     start(root)
#     root.mainloop()



   
# # This function is calling the window
# if __name__ == '__main__':
#     call() 
   
    





